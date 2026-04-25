"""
PII masking engine — builds replacement map and applies to text/docx/xlsx.
"""
import os
from docx import Document
import openpyxl
from openpyxl.styles import PatternFill

from config import LLM_CHUNK
from modules.pii.patterns import MASK_LABELS
from modules.pii.detector import detect_pii_with_regex, detect_pii_with_llm, is_safe_to_mask


def build_replacement_map(full_text: str, model) -> tuple:
    """
    Core engine: regex (instant) + optional AI (max 3 chunks × 15s = 45s cap).
    Regex always returns. AI is a bonus — any crash/timeout is silently ignored.
    Bug Fix 6: replacement map sorted by length desc — longer strings replaced
    first so partial matches don't fragment longer PII values.
    """
    detections = []

    # STEP 1: Regex — always runs, instant
    for entity in detect_pii_with_regex(full_text):
        label = MASK_LABELS.get(entity["type"], f"[{entity['type'].upper()}]")
        detections.append({"original": entity["text"], "masked": label,
                           "type": entity["type"], "source": "Regex"})

    # STEP 2: AI — max 3 chunks of LLM_CHUNK chars, 15s timeout each
    if model:
        sample = full_text[:1500]
        chunks = [sample[i:i + LLM_CHUNK] for i in range(0, len(sample), LLM_CHUNK)]
        for chunk in chunks[:3]:
            if not chunk.strip():
                continue
            for entity in detect_pii_with_llm(chunk, model):
                label = MASK_LABELS.get(entity["type"], f"[{entity['type'].upper()}]")
                detections.append({"original": entity["text"], "masked": label,
                                   "type": entity["type"], "source": "AI"})

    # STEP 3: Build replacement map — sorted by length descending (Bug Fix 6)
    replacement_map = {}
    seen = set()
    for entity in sorted(detections, key=lambda x: len(x["original"]), reverse=True):
        orig = entity["original"].strip()
        if not orig or orig in seen or len(orig) < 2:
            continue
        if orig.startswith('[') and orig.endswith(']'):
            continue
        if not is_safe_to_mask(orig, entity["type"]):
            continue
        seen.add(orig)
        replacement_map[orig] = entity["masked"]

    return replacement_map, detections


def apply_replacements(text: str, replacement_map: dict) -> str:
    """Apply all replacements. Longer strings applied first."""
    for orig in sorted(replacement_map.keys(), key=len, reverse=True):
        text = text.replace(orig, replacement_map[orig])
    return text


# ── File processors ───────────────────────────────────────────────────────────

def process_txt(file_path: str, model) -> tuple:
    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
        content = f.read()
    replacement_map, detections = build_replacement_map(content, model)
    masked = apply_replacements(content, replacement_map)
    out_path = file_path + "_masked.txt"
    with open(out_path, 'w', encoding='utf-8') as f:
        f.write(masked)
    return out_path, detections, masked[:500]


def process_docx(file_path: str, model) -> tuple:
    doc = Document(file_path)

    full_text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text_parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        full_text_parts.append(para.text)

    full_text = "\n".join(full_text_parts)
    print(f"[Mask] Building replacement map from {len(full_text_parts)} paragraphs...")
    replacement_map, detections = build_replacement_map(full_text, model)
    print(f"[Mask] Found {len(replacement_map)} unique PII items — applying to document...")

    preview_parts = []
    for para in doc.paragraphs:
        if not para.text.strip():
            continue
        masked_text = apply_replacements(para.text, replacement_map)
        preview_parts.append(masked_text)
        if masked_text != para.text and para.runs:
            for run in para.runs:
                run.text = ""
            para.runs[0].text = masked_text

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        masked_text = apply_replacements(para.text, replacement_map)
                        if masked_text != para.text and para.runs:
                            for run in para.runs:
                                run.text = ""
                            para.runs[0].text = masked_text

    out_path = file_path + "_masked.docx"
    doc.save(out_path)
    return out_path, detections, "\n".join(preview_parts[:20])


def process_xlsx(file_path: str, model) -> tuple:
    wb = openpyxl.load_workbook(file_path)
    yellow_fill = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")

    all_text_parts = []
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if cell.value and isinstance(cell.value, str) and cell.value.strip():
                    all_text_parts.append(cell.value)

    full_text = "\n".join(all_text_parts)
    replacement_map, detections = build_replacement_map(full_text, model)

    preview_parts = []
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if cell.value and isinstance(cell.value, str) and cell.value.strip():
                    masked_text = apply_replacements(cell.value, replacement_map)
                    if masked_text != cell.value:
                        preview_parts.append(f"{cell.coordinate}: {cell.value} → {masked_text}")
                        cell.value = masked_text
                        cell.fill = yellow_fill

    out_path = file_path + "_masked.xlsx"
    wb.save(out_path)
    return out_path, detections, "\n".join(preview_parts[:20])
