"""
PII Masker routes — Blueprint: sanitize UI, mask file, text masking, file download.
"""
import io
import logging
import os
import uuid

from flask import Blueprint, request, jsonify, send_file, send_from_directory

from config import (
    UPLOAD_FOLDER, MAX_FILE_SIZE, rate_limit,
    make_download_token, verify_download_token,
)
from llm.client import get_model
from modules.pii.masker import (
    build_replacement_map, apply_replacements,
    process_txt, process_docx, process_xlsx,
)

logger = logging.getLogger(__name__)
pii_bp = Blueprint('pii', __name__)


@pii_bp.route('/sanitize')
def sanitize():
    return send_from_directory('static', 'sanitize.html')


@pii_bp.route('/api/mask', methods=['POST'])
@rate_limit("10 per minute; 100 per hour")
def mask_file():
    if 'file' not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files['file']
    if not file.filename:
        return jsonify({"error": "Empty filename"}), 400

    file.seek(0, 2)
    size = file.tell()
    file.seek(0)
    if size > MAX_FILE_SIZE:
        return jsonify({"error": "File too large. Max 50MB allowed."}), 400

    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in ['.txt', '.docx', '.xlsx', '.xls', '.csv']:
        return jsonify({"error": f"Unsupported format: {ext}. Supported: txt, docx, xlsx, csv"}), 400

    # UUID prefix prevents filename collision when concurrent users upload same filename
    tmp_name = f"{uuid.uuid4().hex}{ext}"
    tmp_path = os.path.join(UPLOAD_FOLDER, tmp_name)
    file.save(tmp_path)

    model = get_model()
    logger.info("[mask] file=%s model=%s", file.filename, model or "regex-only")

    out_path = None
    try:
        if ext in ['.txt', '.csv']:
            out_path, detections, preview = process_txt(tmp_path, model)
        elif ext == '.docx':
            out_path, detections, preview = process_docx(tmp_path, model)
        elif ext in ['.xlsx', '.xls']:
            out_path, detections, preview = process_xlsx(tmp_path, model)

        seen_orig = set()
        unique_detections = []
        for d in detections:
            if d["original"] not in seen_orig:
                seen_orig.add(d["original"])
                unique_detections.append(d)

        out_basename = os.path.basename(out_path)
        return jsonify({
            "success":           True,
            "original_filename": file.filename,
            "masked_filename":   out_basename,
            "detections":        unique_detections,
            "detection_count":   len(unique_detections),
            "preview":           preview[:500],
            "download_token":    make_download_token(out_basename),
            "ai_used":           model is not None,
            "model":             model or "regex-only",
        })
    except Exception as e:
        logger.error("[mask] %s", e, exc_info=True)
        return jsonify({"error": "Internal server error"}), 500
    finally:
        # Always clean up the uploaded input file regardless of success/failure
        try:
            os.unlink(tmp_path)
        except OSError:
            pass


@pii_bp.route('/api/download/<path:token>')
def download_file(token):
    safe_name = verify_download_token(token)
    if not safe_name:
        return jsonify({"error": "Invalid or expired download token"}), 403

    file_path = os.path.join(UPLOAD_FOLDER, safe_name)
    if not os.path.exists(file_path):
        return jsonify({"error": "File not found or already downloaded"}), 404

    response = send_file(file_path, as_attachment=True, download_name=safe_name)
    # Clean up after serving — prevent disk accumulation on long-running pod
    # On Unix, unlink while fd is open is safe: data remains until fd closes.
    try:
        os.unlink(file_path)
    except OSError:
        pass
    return response


@pii_bp.route('/api/extract-text', methods=['POST'])
@rate_limit("20 per minute; 300 per hour")
def extract_text_api():
    """Server-side text extraction for binary formats (docx, pdf, xlsx)."""
    if 'file' not in request.files:
        return jsonify({"error": "No file provided"}), 400
    f    = request.files['file']
    name = (f.filename or '').lower()

    try:
        if name.endswith('.docx'):
            from docx import Document as DocxDocument
            doc   = DocxDocument(io.BytesIO(f.read()))
            parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        parts.append(' | '.join(cells))
            text = '\n'.join(parts)
            if not text.strip():
                return jsonify({"error": "Document appears empty or unreadable"}), 400
            return jsonify({"text": text})

        elif name.endswith('.pdf'):
            data = f.read()
            try:
                import pdfplumber
                with pdfplumber.open(io.BytesIO(data)) as pdf:
                    text = '\n'.join(page.extract_text() or '' for page in pdf.pages)
                return jsonify({"text": text})
            except ImportError:
                pass
            try:
                import pypdf
                reader = pypdf.PdfReader(io.BytesIO(data))
                text   = '\n'.join(page.extract_text() or '' for page in reader.pages)
                return jsonify({"text": text})
            except ImportError:
                return jsonify({"error": "PDF extraction unavailable. Install: pip install pdfplumber"}), 500

        elif name.endswith(('.xlsx', '.xls')):
            try:
                import openpyxl
                wb    = openpyxl.load_workbook(io.BytesIO(f.read()), data_only=True)
                parts = []
                for sheet in wb.worksheets:
                    parts.append(f'[Sheet: {sheet.title}]')
                    for row in sheet.iter_rows(values_only=True):
                        cells = [str(c) for c in row if c is not None]
                        if cells:
                            parts.append(' | '.join(cells))
                return jsonify({"text": '\n'.join(parts)})
            except ImportError:
                return jsonify({"error": "Excel extraction unavailable. Install: pip install openpyxl"}), 500

        else:
            text = f.read().decode('utf-8', errors='replace')
            return jsonify({"text": text})

    except Exception as e:
        logger.error("[extract-text] %s", e, exc_info=True)
        return jsonify({"error": "Internal server error"}), 500


@pii_bp.route('/api/mask-text', methods=['POST'])
@rate_limit("30 per minute; 300 per hour")
def mask_text_api():
    data = request.json
    if not data or 'text' not in data:
        return jsonify({"error": "No text provided"}), 400
    try:
        text = data['text'][:10000]
        model = get_model()
        replacement_map, detections = build_replacement_map(text, model)
        masked = apply_replacements(text, replacement_map)
        return jsonify({
            "masked_text":     masked,
            "detections":      detections,
            "detection_count": len(replacement_map),
        })
    except Exception as e:
        logger.error("[mask-text] %s", e, exc_info=True)
        return jsonify({"error": "Internal server error"}), 500
