"""
Chat routes — Blueprint: streaming AI chat with multi-turn history,
stop generation, edit prompt, session restore, DOCX download.
"""
import json
import tempfile

from flask import Blueprint, request, jsonify, send_file, send_from_directory, Response

from config import USE_CLOUD, OPENROUTER_MODEL, UPLOAD_FOLDER, rate_limit, _LLM_SEM
from llm.client import _stream_llama_server, _stream_openrouter, get_model, LLM_MODEL
from modules.chat.session import _CHAT_SESSIONS, _CHAT_SYSTEM_PROMPT

chat_bp = Blueprint('chat', __name__)


@chat_bp.route('/chat')
def chat_page():
    return send_from_directory('static', 'chat.html')


@chat_bp.route('/api/chat/stream', methods=['POST'])
@rate_limit("20 per minute; 200 per hour")
def chat_stream():
    data       = request.json or {}
    session_id = data.get('session_id', '').strip()
    message    = data.get('message', '').strip()

    if not message:
        return jsonify({"error": "message required"}), 400
    if not session_id:
        return jsonify({"error": "session_id required"}), 400

    if session_id not in _CHAT_SESSIONS:
        _CHAT_SESSIONS[session_id] = []

    _CHAT_SESSIONS[session_id].append({"role": "user", "content": message})
    messages = [{"role": "system", "content": _CHAT_SYSTEM_PROMPT}] + _CHAT_SESSIONS[session_id]
    model    = OPENROUTER_MODEL if USE_CLOUD else (get_model() or LLM_MODEL)

    def generate():
        full_response = []
        try:
            with _LLM_SEM:
                token_gen = (
                    _stream_openrouter(messages)
                    if USE_CLOUD
                    else _stream_llama_server(messages, model)
                )
                for token in token_gen:
                    full_response.append(token)
                    yield f"data: {json.dumps({'token': token})}\n\n"
            # Normal completion — save assistant response
            if full_response:
                _CHAT_SESSIONS[session_id].append(
                    {"role": "assistant", "content": "".join(full_response)}
                )
            else:
                # LLM yielded zero tokens — remove the orphaned user message (Bug 1 fix)
                msgs = _CHAT_SESSIONS.get(session_id, [])
                if msgs and msgs[-1]["role"] == "user":
                    msgs.pop()
            yield "data: [DONE]\n\n"
        except Exception as e:
            print(f"[chat/stream] error: {e}")
            # Remove orphaned user message (no paired assistant response will be saved)
            msgs = _CHAT_SESSIONS.get(session_id, [])
            if msgs and msgs[-1]["role"] == "user":
                msgs.pop()
            yield f"data: {json.dumps({'error': str(e)})}\n\n"
            yield "data: [DONE]\n\n"

    return Response(
        generate(),
        content_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@chat_bp.route('/api/chat/session', methods=['GET'])
def get_chat_session():
    """Return stored messages — used by frontend to restore on page reload."""
    session_id = request.args.get('session_id', '').strip()
    messages   = _CHAT_SESSIONS.get(session_id, [])
    return jsonify({"messages": messages})


@chat_bp.route('/api/chat/truncate', methods=['POST'])
def truncate_chat_session():
    """Keep only the first keep_up_to messages — used by Edit to discard subsequent turns."""
    data       = request.json or {}
    session_id = data.get('session_id', '').strip()
    try:
        keep_up_to = max(0, int(data.get('keep_up_to', 0)))
    except (ValueError, TypeError):
        return jsonify({"error": "keep_up_to must be a non-negative integer"}), 400
    if session_id in _CHAT_SESSIONS:
        _CHAT_SESSIONS[session_id] = _CHAT_SESSIONS[session_id][:keep_up_to]
    return jsonify({"success": True})


@chat_bp.route('/api/chat/stop', methods=['POST'])
def stop_chat_stream():
    """Client-initiated stop: remove orphaned user message if generation was aborted."""
    data       = request.json or {}
    session_id = data.get('session_id', '').strip()
    msgs       = _CHAT_SESSIONS.get(session_id, [])
    if msgs and msgs[-1]["role"] == "user":
        msgs.pop()
    return jsonify({"success": True})


@chat_bp.route('/api/chat/session', methods=['DELETE'])
def clear_chat_session():
    data       = request.json or {}
    session_id = data.get('session_id', '').strip()
    if session_id in _CHAT_SESSIONS:
        del _CHAT_SESSIONS[session_id]
    return jsonify({"success": True})


@chat_bp.route('/api/chat/download-docx', methods=['POST'])
def download_chat_docx():
    from docx import Document
    data       = request.json or {}
    session_id = data.get('session_id', '').strip()
    messages   = _CHAT_SESSIONS.get(session_id, [])

    if not messages:
        return jsonify({"error": "No conversation to download"}), 400

    doc = Document()
    doc.add_heading("Chat Conversation", 0)
    for msg in messages:
        p = doc.add_paragraph()
        p.add_run(f"{msg['role'].capitalize()}: ").bold = True
        p.add_run(msg['content'])
        doc.add_paragraph()

    tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False, dir=UPLOAD_FOLDER)
    doc.save(tmp.name)
    tmp.close()
    return send_file(
        tmp.name,
        as_attachment=True,
        download_name='chat_conversation.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )
