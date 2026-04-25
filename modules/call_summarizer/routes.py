"""
Call Summarizer routes — Blueprint: UI page, transcript analysis, DOCX download.
"""
import re
import tempfile

import requests
from flask import Blueprint, request, jsonify, send_file, send_from_directory

from config import USE_CLOUD, OPENROUTER_MODEL, rate_limit
from llm.client import get_model, _call_llm, _model_timeout
from modules.project_docs.parser import _parse_json_response

call_summarizer_bp = Blueprint('call_summarizer', __name__)


@call_summarizer_bp.route('/call-summarizer')
def call_summarizer():
    return send_from_directory('static', 'call_summarizer.html')


@call_summarizer_bp.route('/api/summarize-call', methods=['POST'])
@rate_limit("20 per minute; 200 per hour")
def summarize_call():
    data   = request.get_json() or {}
    source = data.get('source', '').strip()
    if not source:
        return jsonify({"error": "No transcript provided"}), 400

    model = OPENROUTER_MODEL if USE_CLOUD else get_model()
    if not model:
        return jsonify({"error": "No LLM backend. Set OPENROUTER_API_KEY or run: llama-server"}), 503

    # 2500 chars preserves names, owner assignments, and risk detail from transcript
    src_short = source[:2500].strip()
    prompt = """\
Role: Expert Business Analyst (15y exp).
Task: Analyze the call transcript below and return a structured JSON summary.
Output: ONLY valid JSON. No markdown, no prose.
Structure: {{"key_takeaways":["1-2 sentence insight"],"action_items":[{{"task":"What to do","owner":"Name or role from transcript (TBD if not mentioned)","deadline":"When (TBD if not mentioned)"}}],"decisions":["Decision made"],"risks":["Risk or blocker raised"],"next_steps":["Immediate follow-up"]}}
Constraints:
- key_takeaways: 10-15 items, each a distinct insight from the call
- action_items: extract all explicitly assigned tasks; use exact names from transcript for owner
- decisions: only firm decisions confirmed on the call
- risks: blockers, concerns, or dependencies raised — include specific details from transcript
- next_steps: concrete follow-ups in priority order, max 7
Every item must be traceable to the transcript. No invented content.
Transcript: {src}""".format(src=src_short)

    try:
        t   = _model_timeout(model)
        raw = _call_llm(model, prompt, timeout=t,
                        options_override={"max_tokens": 1200}, format_json=True)
        result = _parse_json_response(raw, 'summary')
        return jsonify({"success": True, "data": result, "model": model})
    except requests.exceptions.Timeout:
        return jsonify({"error": "Timed out. Try a shorter transcript or faster model."}), 504
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@call_summarizer_bp.route('/api/download-call-docx', methods=['POST'])
def download_call_docx():
    from docx import Document
    payload         = request.get_json() or {}
    d               = payload.get('data', {})
    transcript_name = payload.get('transcript_name', '').strip()

    safe_name = re.sub(r'[^\w\s-]', '', transcript_name) if transcript_name else ''
    filename  = f"Call Summary - {safe_name}.docx" if safe_name else "Call Summary.docx"

    doc = Document()
    doc.add_heading(f"Call Summary{' — ' + transcript_name if transcript_name else ''}", 0)

    sections = [
        ('key_takeaways',  '💡 Key Takeaways',    'list_numbered'),
        ('action_items',   '✅ Action Items',      'action'),
        ('decisions',      '🎯 Decisions Made',    'list_bullet'),
        ('risks',          '⚠️  Risks & Blockers', 'list_bullet'),
        ('next_steps',     '➡️  Next Steps',       'list_numbered'),
    ]

    for field, heading, style in sections:
        doc.add_heading(heading, level=1)
        items = d.get(field, [])
        if not items:
            doc.add_paragraph('None recorded.').italic = True
            continue
        for i, item in enumerate(items):
            if field == 'action_items':
                task     = item.get('task', str(item))     if isinstance(item, dict) else str(item)
                owner    = item.get('owner', 'TBD')        if isinstance(item, dict) else 'TBD'
                deadline = item.get('deadline', 'TBD')     if isinstance(item, dict) else 'TBD'
                p = doc.add_paragraph()
                p.add_run(f"{i + 1}. {task}").bold = True
                p.add_run(f"\n   Owner: {owner}  |  Deadline: {deadline}")
            elif style == 'list_numbered':
                doc.add_paragraph(f"{i + 1}. {item}")
            else:
                doc.add_paragraph(f"• {item}")

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(tmp.name)
    tmp.close()
    return send_file(
        tmp.name,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )
