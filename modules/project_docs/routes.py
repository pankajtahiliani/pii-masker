"""
Project Documentation routes — Blueprint: UI page, artifact generation, DOCX download.
"""
import json
import re
import tempfile

import requests
from flask import Blueprint, request, jsonify, send_file, send_from_directory

from config import USE_CLOUD, OPENROUTER_MODEL, rate_limit
from llm.client import get_model, _call_llm, _model_profile, _model_timeout
from modules.project_docs.prompts import _ARTIFACT_PROMPTS, _ARTIFACT_PREDICT, _ARTIFACT_ARRAY_KEYS
from modules.project_docs.parser import _parse_json_response

project_docs_bp = Blueprint('project_docs', __name__)


@project_docs_bp.route('/project-docs')
def project_docs():
    return send_from_directory('static', 'project_docs.html')


@project_docs_bp.route('/api/generate-artifact', methods=['POST'])
@rate_limit("30 per minute; 300 per hour")
def generate_artifact():
    """Generate a single Agile artifact. Called N times by the frontend."""
    data           = request.json or {}
    source         = data.get('source', '')
    artifact       = data.get('artifact', '')
    model_override = (data.get('model') or '').strip()

    if not source or not artifact:
        return jsonify({"error": "source and artifact are required"}), 400
    if artifact not in _ARTIFACT_PROMPTS:
        return jsonify({"error": f"Unknown artifact: {artifact}"}), 400

    # Priority: UI override > cloud > auto-pick
    if model_override:
        model = model_override
    else:
        model = OPENROUTER_MODEL if USE_CLOUD else get_model()
    if not model:
        return jsonify({
            "error": "No LLM backend available. Set OPENROUTER_API_KEY env var OR start llama-server"
        }), 503

    base_predict = _ARTIFACT_PREDICT.get(artifact, 1000)
    profile      = _model_profile(model)
    predict      = max(400, int(base_predict * profile["predict_mult"]))

    src_short = source[:3000].strip()
    prompt    = _ARTIFACT_PROMPTS[artifact].format(src=src_short)

    try:
        t = profile["timeout"]
        print(f"[artifact] key={artifact} model={model} src_len={len(src_short)} max_tokens={predict} timeout={t}s")
        raw    = _call_llm(model, prompt, timeout=t,
                           options_override={"max_tokens": predict}, format_json=True)
        result = _parse_json_response(raw, artifact)
        print(f"[artifact] key={artifact} done, result_type={type(result).__name__}")
        return jsonify({"success": True, "artifact": artifact, "data": result, "model": model})
    except requests.exceptions.Timeout:
        return jsonify({"error": f"Timed out generating {artifact}. Try a faster model."}), 504
    except Exception as e:
        print(f"[artifact error] {artifact}: {e}")
        return jsonify({"error": str(e)}), 500


@project_docs_bp.route('/api/refine-artifact', methods=['POST'])
def refine_artifact():
    """Re-generate an artifact using the existing output as context for improvement."""
    data      = request.json or {}
    artifact  = data.get('artifact', '')
    source    = data.get('source', '')
    current   = data.get('current_output', '')
    feedback  = data.get('feedback', '')

    if not artifact or not source:
        return jsonify({"error": "artifact and source required"}), 400
    if artifact not in _ARTIFACT_PROMPTS:
        return jsonify({"error": f"Unknown artifact: {artifact}"}), 400

    model = OPENROUTER_MODEL if USE_CLOUD else get_model()
    if not model:
        return jsonify({"error": "No LLM backend. Set OPENROUTER_API_KEY or start llama-server"}), 503

    predict       = _ARTIFACT_PREDICT.get(artifact, 1000)
    src_short     = source[:3000].strip()
    artifact_labels = {
        'backlog': 'Product Backlog', 'sprint_plan': 'Sprint Planning',
        'sprint_review': 'Sprint Reviews', 'retrospective': 'Retrospective',
        'risk_register': 'Risk Register', 'test_cases': 'Test Cases',
    }
    label           = artifact_labels.get(artifact, artifact.replace('_', ' ').title())
    current_summary = current[:800].strip() if current else ''
    feedback_line   = f'\nUser feedback: {feedback.strip()}' if feedback and feedback.strip() else ''

    refine_prompt = (
        f"Review the previously generated {label} shown below and produce a MORE DETAILED, "
        f"COMPREHENSIVE version with additional relevant entries specific to the project."
        f"{feedback_line}\n"
        f"Previous output (partial): {current_summary}\n\n"
        f"Now generate a complete, improved {label}.\n\n"
        + _ARTIFACT_PROMPTS[artifact].format(src=src_short)
    )

    try:
        t   = _model_timeout(model)
        print(f"[refine] key={artifact} model={model} predict={predict} timeout={t}s")
        raw    = _call_llm(model, refine_prompt, timeout=t,
                           options_override={"max_tokens": predict}, format_json=True)
        result = _parse_json_response(raw, artifact)
        return jsonify({"success": True, "artifact": artifact, "data": result, "model": model})
    except requests.exceptions.Timeout:
        return jsonify({"error": "Timed out. Try a faster model."}), 504
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@project_docs_bp.route('/api/generate-docs', methods=['POST'])
@rate_limit("5 per minute; 60 per hour")
def generate_docs():
    """Kept for backwards compat — calls generate-artifact N× internally."""
    data      = request.json or {}
    source    = data.get('source', '')
    artifacts = data.get('artifacts', list(_ARTIFACT_PROMPTS.keys()))

    model = OPENROUTER_MODEL if USE_CLOUD else get_model()
    if not model:
        return jsonify({"error": "No LLM backend. Set OPENROUTER_API_KEY or start llama-server"}), 503

    documents = {}
    for key in artifacts:
        if key not in _ARTIFACT_PROMPTS:
            continue
        src_short = source[:1500].strip()
        prompt    = _ARTIFACT_PROMPTS[key].format(src=src_short)
        try:
            raw  = _call_llm(model, prompt, timeout=90)
            documents[key] = _parse_json_response(raw, key)
        except Exception as e:
            print(f"[generate-docs] {key} failed: {e}")
            documents[key] = {"_raw": str(e)}

    return jsonify({"success": True, "documents": documents, "model": model})


@project_docs_bp.route('/api/download-docx', methods=['POST'])
def download_docx():
    from docx import Document
    payload        = request.get_json() or {}
    data           = payload.get('data', payload)
    project_name   = payload.get('project_name', '').strip()
    artifact_label = payload.get('artifact_label', '').strip()

    safe_artifact = re.sub(r'[^\w\s-]', '', artifact_label or 'Agile Project Documents')
    safe_project  = re.sub(r'[^\w\s-]', '', project_name) if project_name else ''
    filename = f"{safe_artifact} - {safe_project}.docx" if safe_project else f"{safe_artifact}.docx"

    doc = Document()
    doc.add_heading(artifact_label or 'Agile Project Documents', 0)
    if project_name:
        doc.add_paragraph(f"Project: {project_name}").italic = True

    artifact_labels = {
        'backlog': 'Product Backlog', 'sprint_plan': 'Sprint Planning',
        'sprint_review': 'Sprint Reviews', 'retrospective': 'Retrospective',
        'risk_register': 'Risk Register', 'test_cases': 'Test Cases',
    }

    for key, artifact_data in data.items():
        label = artifact_labels.get(key, key.replace('_', ' ').title())
        doc.add_heading(label, level=1)

        if key == 'backlog':
            if isinstance(artifact_data, list):
                epics = artifact_data
            elif isinstance(artifact_data, dict) and '_raw' in artifact_data:
                try:
                    p = json.loads(artifact_data['_raw'])
                    epics = p if isinstance(p, list) else p.get('epics', [])
                except Exception:
                    epics = []
            else:
                epics = artifact_data.get('epics', []) if isinstance(artifact_data, dict) else []
            for epic in epics:
                doc.add_heading(f"{epic.get('id','')} {epic.get('title','Epic')}".strip(), level=2)
                if epic.get('description'):
                    doc.add_paragraph(epic['description'])
                for s in epic.get('stories', []):
                    pts = s.get('story_points', s.get('points', '?'))
                    pri = s.get('priority', '')
                    p = doc.add_paragraph()
                    p.add_run(f"{s.get('id','')} {s.get('title','')}").bold = True
                    p.add_run(f"  [{pts} SP]  [{pri}]")
                    if s.get('description'):
                        doc.add_paragraph(s['description'], style='List Bullet')
                    for ac in s.get('acceptance_criteria', []):
                        doc.add_paragraph(f"✓ {ac}", style='List Bullet')

        elif key == 'sprint_plan':
            if isinstance(artifact_data, list):
                sprints = artifact_data
            elif isinstance(artifact_data, dict) and '_raw' in artifact_data:
                try:
                    p = json.loads(artifact_data['_raw'])
                    sprints = p if isinstance(p, list) else p.get('sprints', [])
                except Exception:
                    sprints = []
            else:
                sprints = artifact_data.get('sprints', []) if isinstance(artifact_data, dict) else []
            for s in sprints:
                doc.add_heading(f"Sprint {s.get('number','')}".strip(), level=2)
                if s.get('goal'):
                    doc.add_paragraph(f"Goal: {s['goal']}")
                doc.add_paragraph(f"Story Points: {s.get('total_points', s.get('points', 0))}")
                stories = s.get('stories', [])
                if stories:
                    ids = [st if isinstance(st, str) else st.get('id', str(st)) for st in stories]
                    doc.add_paragraph("Stories: " + ", ".join(ids))
                if s.get('deliverable'):
                    doc.add_paragraph(f"Deliverable: {s['deliverable']}")

        elif key == 'sprint_review':
            reviews = artifact_data if isinstance(artifact_data, list) else artifact_data.get('reviews', [])
            for r in reviews:
                doc.add_heading(f"Sprint {r.get('sprint','')} Review".strip(), level=2)
                if r.get('planned_stories'):
                    doc.add_paragraph("Planned: " + ", ".join(str(s) for s in r['planned_stories']))
                if r.get('completed_stories'):
                    doc.add_paragraph("Completed: " + ", ".join(str(s) for s in r['completed_stories']))
                if r.get('demo_notes'):
                    doc.add_paragraph(f"Demo Notes: {r['demo_notes']}")
                if r.get('feedback_template'):
                    doc.add_paragraph(f"Stakeholder Feedback: {r['feedback_template']}")
                if r.get('next_sprint_adjustments'):
                    doc.add_paragraph(f"Next Sprint Adjustments: {r['next_sprint_adjustments']}")

        elif key == 'retrospective':
            for field, section_label in [
                ('went_well', 'What Went Well'), ('improve', 'What to Improve'),
                ('action_items', 'Action Items'), ('team_health_check', 'Team Health Check'),
            ]:
                items = artifact_data.get(field, []) if isinstance(artifact_data, dict) else []
                if items:
                    doc.add_heading(section_label, level=2)
                    for item in items:
                        doc.add_paragraph(f"• {item}")

        elif key == 'risk_register':
            if isinstance(artifact_data, dict) and '_raw' in artifact_data:
                try:
                    parsed = json.loads(artifact_data['_raw'])
                    artifact_data = parsed if isinstance(parsed, list) else parsed.get('risks', parsed)
                except Exception:
                    pass
            risks = artifact_data if isinstance(artifact_data, list) else artifact_data.get('risks', [])
            for r in risks:
                title = f"{r.get('id','')} {r.get('title', r.get('risk','Risk'))}".strip()
                doc.add_heading(title, level=2)
                if r.get('description'):
                    doc.add_paragraph(r['description'])
                doc.add_paragraph(
                    f"Probability: {r.get('probability','M')}  |  "
                    f"Impact: {r.get('impact','M')}  |  "
                    f"Severity: {r.get('severity','Medium')}")
                if r.get('mitigation'):
                    doc.add_paragraph(f"Mitigation: {r['mitigation']}")

        elif key == 'test_cases':
            cases = artifact_data if isinstance(artifact_data, list) else artifact_data.get('test_cases', [])
            for tc in cases:
                doc.add_heading(f"{tc.get('id','')} {tc.get('title','Test Case')}".strip(), level=2)
                doc.add_paragraph(
                    f"Type: {tc.get('type','')}  |  "
                    f"Story: {tc.get('story_id','')}  |  "
                    f"Priority: {tc.get('priority','')}")
                if tc.get('precondition'):
                    doc.add_paragraph(f"Precondition: {tc['precondition']}")
                for i, step in enumerate(tc.get('steps', []), 1):
                    doc.add_paragraph(f"{i}. {step}", style='List Number')
                if tc.get('expected_result'):
                    doc.add_paragraph(f"Expected Result: {tc['expected_result']}")

        else:
            txt = artifact_data if isinstance(artifact_data, str) else json.dumps(artifact_data, indent=2)
            doc.add_paragraph(txt)

        doc.add_page_break()

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(tmp.name)
    tmp.close()
    return send_file(
        tmp.name,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )
