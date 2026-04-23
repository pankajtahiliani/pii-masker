"""
PII Masker - Sensitive Data Masking Tool
Uses local llama-server (llama.cpp) for AI-powered PII detection
Works completely offline

v6 - Bug fixes applied after QC review:
   Fix 1: "Email"/"Date" column headers no longer masked as [PASSWORD]
   Fix 2: Date strings (01-Jan-2025) no longer match password pattern
   Fix 3: Phone numbers no longer split into [ZIP_CODE][ZIP_CODE]
   Fix 4: Single-char middle segment usernames (anil_v_arch) now masked
   Fix 5: [LEGAL_ACT] label replaced with proper [REDACTED] in MASK_LABELS
   Fix 6: API keys/tokens no longer partially broken by ZIP pattern
   All 8-continent PII patterns preserved unchanged
"""

import os
import re
import json
import uuid
import tempfile
import threading
import requests
from flask import Flask, request, jsonify, send_file, send_from_directory, Response
from flask_cors import CORS
from docx import Document
import openpyxl
from openpyxl.styles import PatternFill

# Optional: flask-limiter for per-IP rate limiting (graceful degrade if absent)
try:
    from flask_limiter import Limiter
    from flask_limiter.util import get_remote_address
    _HAS_LIMITER = True
except ImportError:
    _HAS_LIMITER = False

app = Flask(__name__, static_folder='static')
CORS(app)

# ── Concurrency guardrails (multi-user deploy) ──────────────────────────────────
# Semaphore caps concurrent LLM calls to match llama-server --parallel slots.
# Set LLM_MAX_CONCURRENT = --parallel value on the server (default 4 Mac, 8 GPU).
_LLM_MAX_CONCURRENT = int(os.environ.get("LLM_MAX_CONCURRENT", "1"))  # dev default; prod: set LLM_MAX_CONCURRENT=8
_LLM_SEM = threading.BoundedSemaphore(_LLM_MAX_CONCURRENT)

# Persistent HTTP connection pool to llama-server — avoids TCP handshake on every request.
# pool_maxsize = parallel slots + headroom for non-streaming calls running alongside streams.
_LLAMA_SESSION = requests.Session()
_LLAMA_SESSION.mount('http://', requests.adapters.HTTPAdapter(
    pool_connections=4, pool_maxsize=_LLM_MAX_CONCURRENT + 2, max_retries=0,
))
_LLAMA_SESSION.mount('https://', requests.adapters.HTTPAdapter(
    pool_connections=2, pool_maxsize=4, max_retries=0,
))

# Rate limiter: 12 artifact-gens/min/IP (1 full doc-set), 200/hr/IP
if _HAS_LIMITER:
    limiter = Limiter(
        get_remote_address,
        app=app,
        default_limits=["120 per minute", "2000 per hour"],
        storage_uri=os.environ.get("LIMITER_STORAGE", "memory://"),
    )
else:
    limiter = None
    print("⚠️  flask-limiter not installed — rate limiting DISABLED. Run: pip install flask-limiter")


def rate_limit(spec):
    """No-op when flask-limiter absent; real limit when present."""
    if _HAS_LIMITER:
        return limiter.limit(spec)
    return lambda f: f

# llama-server (llama.cpp) — OpenAI-compatible endpoint
_LLAMA_BASE     = os.environ.get("LLAMA_SERVER_URL", "http://localhost:8080")
LLAMA_CHAT_URL  = f"{_LLAMA_BASE}/v1/chat/completions"
LLAMA_MODELS_URL = f"{_LLAMA_BASE}/v1/models"
LLAMA_HEALTH_URL = f"{_LLAMA_BASE}/health"

# ── Optional cloud inference (OpenRouter) ───────────────────────────────────────
# LOCAL-ONLY by default. Privacy = value prop. Cloud would send PII off-device.
# Opt-in only for users who explicitly want speed over privacy:
#   export PII_CLOUD_OPT_IN=1 OPENROUTER_API_KEY=sk-or-v1-...
# Without both env vars set, app stays fully local via llama.cpp.
_CLOUD_OPT_IN        = os.environ.get("PII_CLOUD_OPT_IN", "").strip() == "1"
OPENROUTER_API_KEY   = os.environ.get("OPENROUTER_API_KEY", "").strip() if _CLOUD_OPT_IN else ""
OPENROUTER_URL       = "https://openrouter.ai/api/v1/chat/completions"
OPENROUTER_MODEL     = os.environ.get("OPENROUTER_MODEL", "google/gemma-3-27b-it:free")
USE_CLOUD            = bool(OPENROUTER_API_KEY)
LLM_TIMEOUT     = 15        # hard timeout per chunk — 15s plenty with num_ctx=512
LLM_CHUNK       = 500       # Reduced 1500→500: smaller prompt = ~3x faster response
MAX_FILE_SIZE   = 50 * 1024 * 1024
# Fixed shared upload dir — all gunicorn threads use same path, no cross-worker miss.
# Temp dir per-process would break download when worker that saved != worker that serves.
UPLOAD_FOLDER   = os.environ.get("UPLOAD_FOLDER", "/tmp/pii_uploads")
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# ── Chat module ────────────────────────────────────────────────────────────────
# In-memory conversation store: {session_id: [{role, content}, ...]}
# Lost on server restart — intentional (privacy, no persistence needed).
_CHAT_SESSIONS: dict = {}

_CHAT_SYSTEM_PROMPT = (
    "You are a helpful AI assistant for software teams. "
    "Help with requirements analysis, effort estimation, user story writing, "
    "risk analysis, Q&A, and any other software-related needs. "
    "Be concise, structured, and practical in your responses."
)

# Single model alias — must match --alias set at llama-server startup.
# Override via LLM_MODEL env var (e.g. LLM_MODEL=gemma3 for speed testing).
LLM_MODEL = os.environ.get("LLM_MODEL", "qwen2.5")

# Per-request inference options for llama-server (OpenAI-compatible fields).
# Context size, GPU layers, threads are startup flags — NOT per-request.
LLM_OPTIONS = {
    "max_tokens":     150,   # NER needs very few output tokens
    "temperature":    0,     # deterministic — no sampling overhead
    "top_k":          1,     # greedy decode — fastest sampling
    "top_p":          1.0,   # disable nucleus sampling overhead
    "repeat_penalty": 1.0,   # disable repetition check overhead
}

# ═══════════════════════════════════════════════════════════════════════════════
# GLOBAL PII PATTERN ENGINE — covers IN, AU, UK, US, EU, CA, SG, UAE + generic
# ═══════════════════════════════════════════════════════════════════════════════

REGEX_PATTERNS = {

    # ── Universal ────────────────────────────────────────────────────────────
    "email":        r'\b[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Z|a-z]{2,}\b',
    "url":          r'https?://[^\s<>"]+|www\.[^\s<>"]+',
    "ip_address":   r'\b(?:\d{1,3}\.){3}\d{1,3}\b',
    "ipv6":         r'\b([0-9a-fA-F]{1,4}:){7}[0-9a-fA-F]{1,4}\b',
    "credit_card":  r'\b(?:\d{4}[\s\-]?){3}\d{4}\b',
    "gender":       r'\b(?:Male|Female|Other|Non-binary|Transgender|Prefer not to say)\b',
    "dob":          r'\b(?:\d{1,2}[\/\-\.]\d{1,2}[\/\-\.]\d{2,4}|\d{4}[\/\-\.]\d{1,2}[\/\-\.]\d{1,2})\b',

    # Strong password pattern (mixed case + digit + symbol, 8+ chars)
    "password":     r'\b(?=[A-Za-z0-9!@#$%^&*()\-_=+{};:,<.>]*[A-Z])(?=[A-Za-z0-9!@#$%^&*()\-_=+{};:,<.>]*[a-z])(?=[A-Za-z0-9!@#$%^&*()\-_=+{};:,<.>]*\d)(?=[A-Za-z0-9!@#$%^&*()\-_=+{};:,<.>]*[!@#$%^&*()\-_=+{};:,<.>])[A-Za-z0-9!@#$%^&*()\-_=+{};:,<.>]{8,}\b',

    # Tech credentials
    "aws_key":      r'\b(?:AKIA|ASIA|AROA)[A-Z0-9]{16}\b',
    "aws_secret":   r'\b[A-Za-z0-9/+=]{40}\b',
    "api_key":      r'\b(?:SG\.|rzp_(?:test|live)_|sk_(?:test|live)_|pk_(?:test|live)_|AC[a-f0-9]{30,}|AIzaSy|gh[pousr]_)[A-Za-z0-9\._\-]{10,}\b',
    "jwt_token":    r'\beyJ[A-Za-z0-9\-_]+\.[A-Za-z0-9\-_]+\.[A-Za-z0-9\-_]+\b',

    # System usernames
    "admin_user":   r'\b(?:admin|prod|uat|dev|root|superuser|sysadmin)_[a-z0-9_]+\b',
    "username":     r'\b[a-z][a-z0-9]*_[a-z][a-z0-9]*(?:_[a-z0-9]+)?\b',  # Bug Fix 4: single-char segments

    # ── INDIA ────────────────────────────────────────────────────────────────
    "in_phone":     r'(?:\+91|0091|91)[\s\-]?[6-9]\d{4}[\s\-]?\d{5}\b',  # Bug Fix 3: allow space in middle
    "in_aadhaar":   r'\b\d{4}[\s\-]\d{4}[\s\-]\d{4}\b',
    "in_pan":       r'\b[A-Z]{5}[0-9]{4}[A-Z]\b',
    "in_gst":       r'\b\d{2}[A-Z]{5}\d{4}[A-Z]\d[Z][A-Z0-9]\b',
    "in_passport":  r'\b[A-Z][1-9]\d{7}\b',
    "in_voter_id":  r'\b[A-Z]{3}[0-9]{7}\b',
    "in_driving":   r'\b[A-Z]{2}\d{2}\s?\d{11}\b',
    "in_pincode":   r'\b[1-9][0-9]{5}\b',
    "in_ifsc":      r'\b[A-Z]{4}0[A-Z0-9]{6}\b',
    "in_upi":       r'\b[\w.\-]+@(?:okicici|oksbi|okaxis|okhdfcbank|paytm|ybl|ibl|axl|upi)\b',

    # ── AUSTRALIA ────────────────────────────────────────────────────────────
    "au_phone":     r'(?:\+61|0061)?[\s\-]?(?:0?[2378]\d{8}|4\d{8})\b',
    "au_tfn":       r'\b\d{3}[\s\-]?\d{3}[\s\-]?\d{3}\b',          # Tax File Number
    "au_abn":       r'\b\d{2}[\s]?\d{3}[\s]?\d{3}[\s]?\d{3}\b',    # ABN
    "au_acn":       r'\b\d{3}[\s]?\d{3}[\s]?\d{3}\b',              # ACN
    "au_medicare":  r'\b\d{4}[\s]?\d{5}[\s]?\d{1}\b',              # Medicare
    "au_postcode":  r'\b(?:0[89]\d{2}|[1-9]\d{3})\b',
    "au_passport":  r'\b[A-Z]{1,2}\d{7}\b',
    "au_drivers":   r'\b\d{8,9}\b',                                 # State licence numbers

    # ── UNITED KINGDOM ───────────────────────────────────────────────────────
    "uk_phone":     r'(?:\+44|0044)?[\s\-]?(?:0?[1-9]\d{9}|7\d{9})\b',
    "uk_nino":      r'\b[A-CEGHJ-PR-TW-Z]{2}\d{6}[A-D]\b',         # National Insurance
    "uk_postcode":  r'\b[A-Z]{1,2}\d[A-Z\d]?\s?\d[A-Z]{2}\b',
    "uk_passport":  r'\b\d{9}\b',
    "uk_sort_code": r'\b\d{2}[\s\-]\d{2}[\s\-]\d{2}\b',
    "uk_bank_acct": r'\b\d{8}\b',
    "uk_nhs":       r'\b\d{3}[\s\-]\d{3}[\s\-]\d{4}\b',            # NHS Number
    "uk_utr":       r'\b\d{10}\b',                                   # UTR (tax)
    "uk_company":   r'\b(?:SC|NI|OC|SO|NC)?\d{6,8}\b',             # Companies House

    # ── UNITED STATES ────────────────────────────────────────────────────────
    "us_phone":     r'(?:\+1[\s\-]?)?(?:\(\d{3}\)[\s\-]?|\d{3}[\s\-])\d{3}[\s\-]\d{4}\b',
    "us_ssn":       r'\b(?!000|666|9\d{2})\d{3}[\s\-]\d{2}[\s\-]\d{4}\b',
    "us_zip":       r'\b\d{5}(?:\-\d{4})?\b',
    "us_passport":  r'\b[A-Z]\d{8}\b',
    "us_ein":       r'\b\d{2}\-\d{7}\b',                           # Employer ID
    "us_itin":      r'\b9\d{2}[\s\-]\d{2}[\s\-]\d{4}\b',          # ITIN
    "us_drivers":   r'\b[A-Z]\d{7}\b',                             # Common format
    "us_dea":       r'\b[A-Z]{2}\d{7}\b',                          # DEA Number
    "us_npi":       r'\b\d{10}\b',                                  # NPI (healthcare)

    # ── CANADA ───────────────────────────────────────────────────────────────
    "ca_phone":     r'(?:\+1[\s\-]?)?(?:\(\d{3}\)[\s\-]?|\d{3}[\s\-])\d{3}[\s\-]\d{4}\b',
    "ca_sin":       r'\b\d{3}[\s\-]\d{3}[\s\-]\d{3}\b',           # SIN
    "ca_postcode":  r'\b[A-Z]\d[A-Z][\s]?\d[A-Z]\d\b',
    "ca_passport":  r'\b[A-Z]{2}\d{6}\b',
    "ca_health":    r'\b\d{10}\b',                                  # Provincial health card

    # ── EUROPEAN UNION (generic + major countries) ────────────────────────────
    "eu_phone":     r'(?:\+(?:33|49|34|39|31|32|41|43|45|46|47|48|351|353|358|420|421|36|40|359|370|371|372|386|385|356|357)[\s\-]?)[\d\s\-\(\)]{7,15}\b',
    "eu_vat":       r'\b(?:GB|FR|DE|IT|ES|NL|BE|PL|SE|AT|DK|FI|PT|IE|CZ|SK|HU|RO|BG|HR|LT|LV|EE|SI|CY|MT|LU)\s?[A-Z0-9]{8,12}\b',
    "de_id":        r'\b[A-Z0-9]{9}\b',                            # German ID
    "fr_nir":       r'\b[12]\s?\d{2}\s?\d{2}\s?\d{2}\s?\d{3}\s?\d{3}\s?\d{2}\b',  # INSEE
    "eu_iban":      r'\b[A-Z]{2}\d{2}[A-Z0-9]{4}\d{7}(?:[A-Z0-9]{0,16})\b',

    # ── UAE / GCC ─────────────────────────────────────────────────────────────
    "ae_phone":     r'(?:\+971|00971)?[\s\-]?(?:0?[2-9]\d{7}|5[024568]\d{7})\b',
    "ae_eid":       r'\b784[\s\-]?\d{4}[\s\-]?\d{7}[\s\-]?\d{1}\b',  # Emirates ID
    "ae_trade_lic": r'\b(?:CN|BN|LN|TL)-?\d{4,10}\b',
    "sa_phone":     r'(?:\+966|00966)?[\s\-]?0?[15]\d{8}\b',
    "sa_iqama":     r'\b[12]\d{9}\b',                              # Saudi Iqama

    # ── SINGAPORE ────────────────────────────────────────────────────────────
    "sg_phone":     r'(?:\+65|0065)?[\s\-]?[689]\d{7}\b',
    "sg_nric":      r'\b[STFGM]\d{7}[A-Z]\b',                     # NRIC/FIN
    "sg_uen":       r'\b\d{9}[A-Z]\b',                             # UEN
    "sg_passport":  r'\bE\d{7}[A-Z]\b',

    # ── ADDRESSES (global structure) ──────────────────────────────────────────
    # Unit/Street number + street type keywords
    "address_unit": r'\b(?:Apt|Apartment|Suite|Ste|Unit|Flat|Floor|Fl|Room|Rm|Level|Lvl)\.?\s*#?\d+[A-Za-z]?\b',
    "address_street": r'\b\d+\s+[A-Za-z0-9\s]{3,30}(?:Street|St|Avenue|Ave|Road|Rd|Boulevard|Blvd|Lane|Ln|Drive|Dr|Court|Ct|Place|Pl|Way|Terrace|Terr|Close|Crescent|Cres|Highway|Hwy|Parade|Circuit|Grove|Gardens?)\b',
    # PO Box
    "po_box":       r'\b(?:P\.?O\.?|Post Office)\s*Box\s*\d+\b',
}

# ── Context-aware extraction (keyword: value — works across all regions) ───────
CONTEXT_LABELS = {
    r'(?i)(?:full\s*(?:legal\s*)?name|customer\s*name|contact\s*name|first\s*name|last\s*name|given\s*name|surname|family\s*name)\s*[:\-–]\s*': (r'[^\n,;]{2,60}', 'full_name'),
    r'(?i)(?:gender|sex)\s*[:\-–]\s*':                                   (r'\S+', 'gender'),
    r'(?i)(?:date\s*of\s*birth|dob|birth\s*date)\s*[:\-–]\s*':          (r'[\d\/\-\.A-Za-z\s]{4,20}', 'dob'),
    r'(?i)(?:address|addr|residence|home\s*address|mailing\s*address|street\s*address)\s*[:\-–]\s*': (r'[^\n]{5,120}', 'address'),
    r'(?i)(?:postcode|postal\s*code|zip\s*code?|pin\s*code)\s*[:\-–]\s*': (r'[A-Z0-9\s\-]{3,10}', 'postcode'),
    r'(?i)(?:username|user\s*name|login\s*(?:id|name)|user\s*id|account\s*name)\s*[:\-–]\s*': (r'\S+', 'username'),
    r'(?i)(?:password|passwd|default\s*password|pwd|passphrase)\s*[:\-–]?\s*': (r'\S{4,}', 'password'),
    r'(?i)(?:mobile|cell(?:phone)?|phone|telephone|tel|contact\s*(?:no|number))\s*[:\-–]\s*': (r'[\+\d\s\-\(\)]{7,20}', 'phone'),
    r'(?i)(?:national\s*id|nat\.?\s*id|id\s*(?:number|no)|passport\s*(?:no|number)|driving\s*licen[sc]e)\s*[:\-–]\s*': (r'[A-Z0-9\s\-]{5,20}', 'national_id'),
    r'(?i)(?:ssn|social\s*security)\s*[:\-–]\s*':                        (r'[\d\s\-]{9,11}', 'us_ssn'),
    r'(?i)(?:aadhaar|aadhar|uid)\s*[:\-–]\s*':                          (r'[\d\s\-]{12,14}', 'in_aadhaar'),
    r'(?i)(?:pan\s*(?:number|no)?)\s*[:\-–]\s*':                        (r'[A-Z0-9]{10}', 'in_pan'),
    r'(?i)(?:tfn|tax\s*file\s*(?:number|no))\s*[:\-–]\s*':              (r'[\d\s]{9,11}', 'au_tfn'),
    r'(?i)(?:nino|national\s*insurance)\s*[:\-–]\s*':                    (r'[A-Z0-9\s]{9,11}', 'uk_nino'),
    r'(?i)(?:medicare)\s*[:\-–]\s*':                                     (r'[\d\s]{10,12}', 'au_medicare'),
    r'(?i)(?:access\s*key|secret\s*(?:key|access)|api\s*key|api\s*token)\s*[:\-–]\s*': (r'\S{10,}', 'api_key'),
    r'(?i)(?:bank\s*account|account\s*(?:no|number))\s*[:\-–]\s*':      (r'[\d\s\-]{6,20}', 'bank_account'),
    r'(?i)(?:sort\s*code|routing\s*(?:no|number))\s*[:\-–]\s*':         (r'[\d\s\-]{6,11}', 'sort_code'),
    r'(?i)(?:iban)\s*[:\-–]\s*':                                        (r'[A-Z0-9\s]{15,34}', 'eu_iban'),
    r'(?i)(?:nationality|citizenship|country\s*of\s*origin)\s*[:\-–]\s*': (r'[A-Za-z\s]{3,30}', 'nationality'),
    r'(?i)(?:ethnicity|race)\s*[:\-–]\s*':                              (r'[A-Za-z\s\-]{3,30}', 'ethnicity'),
    r'(?i)(?:religion|faith)\s*[:\-–]\s*':                              (r'[A-Za-z\s]{3,20}', 'religion'),
    r'(?i)(?:salary|annual\s*(?:salary|income|ctc)|compensation|remuneration)\s*[:\-–]\s*': (r'[\d,\.\s\$£€₹]{3,20}', 'salary'),
    r'(?i)(?:dbs|criminal\s*record|background\s*check)\s*[:\-–]\s*':    (r'[A-Za-z0-9\s]{2,30}', 'background_check'),
    r'(?i)(?:disability|medical\s*condition|health\s*condition)\s*[:\-–]\s*': (r'[^\n]{3,60}', 'health'),
}

# ── Proper name detection (global — not region-specific) ─────────────────────
NAME_PATTERNS = [
    # Honourifics first (most reliable signal)
    r'\b(?:Mr|Mrs|Ms|Miss|Dr|Prof|Sir|Dame|Rev|Hon|Capt|Maj|Col|Gen|Sgt|Fr|Sr|Br)\.?\s+[A-Z][a-z]+(?:\s+[A-Z][a-z]+){0,3}\b',
    # Three-word names (First Middle Last) — high confidence
    r'\b[A-Z][a-z]{1,20}\s+[A-Z][a-z]{1,20}\s+[A-Z][a-z]{1,20}\b',
    # Two-word names (First Last) — needs NOT_NAMES filter
    r'\b[A-Z][a-z]{1,20}\s+[A-Z][a-z]{1,20}\b',
]

# ── Words that look like names but aren't — expanded for global coverage ───────
NOT_NAMES = {
    # Job titles & roles
    'Business','Analyst','Product','Owner','Technical','Architect','Scrum','Master',
    'DevOps','Engineer','Finance','Director','Manager','Senior','Junior','Lead','Head',
    'Chief','Vice','President','Officer','Associate','Consultant','Specialist','Executive',
    'Administrator','Coordinator','Supervisor','Inspector','Controller','Partner',
    # Document words
    'Software','Requirements','Specification','Document','Control','Overview','Background',
    'Objective','Scope','Summary','Revision','History','Contents','Section','Heading',
    'Appendix','Annex','Version','Status','Draft','Final','Initial','Approved','Signed',
    # Actions & states
    'Access','Level','Default','Password','Username','Login','Logout','Register','Reset',
    'Primary','Secondary','Alternate','Optional','Mandatory','Required','Recommended',
    'Encrypted','Hashed','Masked','Redacted','Restricted','Confidential','Public',
    # Technical
    'Module','System','Platform','Service','Gateway','Bucket','Token','Secret','Key',
    'Environment','Development','Production','Staging','Testing','Integration','Deployment',
    'Backend','Frontend','Database','Server','Client','Infrastructure','Architecture',
    'Application','Interface','Dashboard','Report','Analytics','Pipeline','Workflow',
    # Geography — cities, countries, states (prevent "New South Wales" → name)
    'Australia','United','Kingdom','States','America','Canada','Singapore','Emirates',
    'India','Pakistan','Bangladesh','Sri','Lanka','Nepal','Bhutan','Maldives',
    'New','South','Wales','Victoria','Queensland','Western','Northern','Territory',
    'England','Scotland','Wales','Ireland','Northern','London','Manchester','Sydney',
    'Melbourne','Brisbane','Perth','Adelaide','Auckland','Wellington','Christchurch',
    'Delhi','Mumbai','Bengaluru','Chennai','Hyderabad','Kolkata','Jaipur','Noida','Pune',
    'Dubai','Abu','Dhabi','Riyadh','Doha','Kuwait','Bahrain','Muscat',
    'California','Texas','Florida','York','Jersey','Hampshire','Mexico','Orleans',
    'Pradesh','Rajasthan','Maharashtra','Karnataka','Kerala','Gujarat','Bihar',
    'Uttar','Madhya','Himachal','Jammu','Kashmir','Uttarakhand','Jharkhand',
    # Address words
    'Street','Avenue','Road','Boulevard','Lane','Drive','Court','Place','Way',
    'Terrace','Close','Crescent','Highway','Parade','Circuit','Grove','Gardens',
    'Sector','Floor','Plot','Flat','Apartment','Suite','Block','Wing','Tower',
    'Phase','Nagar','Marg','Chowk','Colony','Layout','Extension','Enclave',
    # Time
    'January','February','March','April','June','July','August','September',
    'October','November','December','Monday','Tuesday','Wednesday','Thursday',
    'Friday','Saturday','Sunday','Today','Tomorrow','Yesterday',
    # Miscellaneous
    'Customer','Employee','Team','Register','Legal','Billing','Contract','Reference',
    'Project','Platform','Services','Super','Admin','Backend','Frontend',
    'Third','Party','Payment','Maps','Documents','Sign','Signature','Approval',
    'Review','Change','Full','Name','Mobile','Phone','Role','Gender','Birth',
    'Field','Type','Notes','Format','Generated','Deployment','Token','Service',
    'Male','Female','Other','True','False','None','Null','Yes','No','Not',
    'Available','Applicable','Provided','Contact','Details','Information',
    'Account','Number','Reference','Code','Serial','Batch','Order','Invoice',
}

MASK_LABELS = {
    # Universal
    "email": "[EMAIL]", "url": "[URL]", "ip_address": "[IP]", "ipv6": "[IP]",
    "gender": "[GENDER]", "dob": "[DOB]", "password": "[PASSWORD]",
    "credit_card": "[CREDIT_CARD]", "aws_key": "[AWS_KEY]", "aws_secret": "[AWS_SECRET]",
    "api_key": "[API_KEY]", "jwt_token": "[JWT_TOKEN]",
    "admin_user": "[USERNAME]", "username": "[USERNAME]",
    "full_name": "[FULL_NAME]", "name": "[NAME]", "phone": "[PHONE]",
    "address": "[ADDRESS]", "address_unit": "[ADDRESS]", "address_street": "[ADDRESS]",
    "po_box": "[ADDRESS]", "postcode": "[POSTCODE]", "national_id": "[NATIONAL_ID]",
    "bank_account": "[BANK_ACCOUNT]", "sort_code": "[SORT_CODE]",
    "nationality": "[NATIONALITY]", "ethnicity": "[ETHNICITY]", "religion": "[RELIGION]",
    "salary": "[SALARY]", "background_check": "[BACKGROUND_CHECK]", "health": "[HEALTH]",
    # India
    "in_phone": "[PHONE]", "in_aadhaar": "[AADHAAR]", "in_pan": "[PAN]",
    "in_gst": "[GST]", "in_passport": "[PASSPORT]", "in_voter_id": "[VOTER_ID]",
    "in_driving": "[DRIVING_LICENCE]", "in_pincode": "[POSTCODE]",
    "in_ifsc": "[IFSC_CODE]", "in_upi": "[UPI_ID]",
    # Australia
    "au_phone": "[PHONE]", "au_tfn": "[TAX_FILE_NO]", "au_abn": "[ABN]",
    "au_acn": "[ACN]", "au_medicare": "[MEDICARE]", "au_postcode": "[POSTCODE]",
    "au_passport": "[PASSPORT]", "au_drivers": "[DRIVING_LICENCE]",
    # UK
    "uk_phone": "[PHONE]", "uk_nino": "[NINO]", "uk_postcode": "[POSTCODE]",
    "uk_passport": "[PASSPORT]", "uk_sort_code": "[SORT_CODE]",
    "uk_bank_acct": "[BANK_ACCOUNT]", "uk_nhs": "[NHS_NUMBER]",
    "uk_utr": "[UTR]", "uk_company": "[COMPANY_NO]",
    # US
    "us_phone": "[PHONE]", "us_ssn": "[SSN]", "us_zip": "[ZIP_CODE]",
    "us_passport": "[PASSPORT]", "us_ein": "[EIN]", "us_itin": "[ITIN]",
    "us_drivers": "[DRIVING_LICENCE]", "us_dea": "[DEA_NO]", "us_npi": "[NPI]",
    # Canada
    "ca_phone": "[PHONE]", "ca_sin": "[SIN]", "ca_postcode": "[POSTCODE]",
    "ca_passport": "[PASSPORT]", "ca_health": "[HEALTH_CARD]",
    # EU
    "eu_phone": "[PHONE]", "eu_vat": "[VAT_NO]", "de_id": "[NATIONAL_ID]",
    "fr_nir": "[NATIONAL_ID]", "eu_iban": "[IBAN]",
    # UAE/GCC
    "ae_phone": "[PHONE]", "ae_eid": "[EMIRATES_ID]", "ae_trade_lic": "[TRADE_LICENCE]",
    "sa_phone": "[PHONE]", "sa_iqama": "[IQAMA]",
    # Singapore
    "sg_phone": "[PHONE]", "sg_nric": "[NRIC]", "sg_uen": "[UEN]",
    "sg_passport": "[PASSPORT]",
    # Fix 5: explicit labels for dynamically generated types
    "legal_act": "[REDACTED]", "act": "[REDACTED]",
    "bank_account": "[BANK_ACCOUNT]", "sort_code": "[SORT_CODE]",
    "dob": "[DOB]", "date": "[DATE]",
}




# ─── llama-server helpers ──────────────────────────────────────────────────────

def get_model():
    """Return configured model alias. llama-server loads one model — no discovery needed."""
    return LLM_MODEL


def detect_pii_with_llm(text: str, model: str) -> list:
    """
    Single chunk AI call with hard 15s timeout.
    Returns [] immediately on any error — regex results already saved.
    """
    prompt = (
        "PII detector. Find: full names, usernames, physical addresses, gender. "
        "Return ONLY a JSON array. Each item: "
        '{"text": "exact string", "type": "name|username|address|gender"}\n\n'
        f"TEXT:\n{text}\n\nJSON:"
    )
    try:
        payload = {
            "model":       model,
            "messages":    [{"role": "user", "content": prompt}],
            "max_tokens":  LLM_OPTIONS["max_tokens"],
            "temperature": LLM_OPTIONS["temperature"],
            "top_k":       LLM_OPTIONS["top_k"],
            "top_p":       LLM_OPTIONS["top_p"],
        }
        resp = requests.post(LLAMA_CHAT_URL, json=payload, timeout=LLM_TIMEOUT)
        if resp.status_code == 200:
            raw = resp.json()["choices"][0]["message"]["content"]
            m = re.search(r'\[.*?\]', raw, re.DOTALL)
            if m:
                entities = json.loads(m.group())
                return [e for e in entities
                        if isinstance(e, dict) and "text" in e and "type" in e
                        and len(e["text"].strip()) > 1]
    except requests.exceptions.Timeout:
        print(f"[AI] Timeout after {LLM_TIMEOUT}s — regex results only")
    except Exception as e:
        print(f"[AI] Error: {e} — regex results only")
    return []


def check_backend_available() -> dict:
    """Check llama-server health and return loaded model info."""
    try:
        r = requests.get(LLAMA_HEALTH_URL, timeout=3)
        if r.status_code == 200:
            return {"available": True, "models": [LLM_MODEL]}
    except:
        pass
    return {"available": False, "models": []}


# ─── Fix 1 & 2: Safe value guards ─────────────────────────────────────────────

# Words that are column/field headers — never mask these as PII values
# Bug Fix 1: "Email", "Date", "Password" as standalone column headers were
# being incorrectly masked by the password pattern or context labels
HEADER_WORDS = {
    'Email', 'Date', 'Password', 'Name', 'Phone', 'Mobile', 'Address',
    'Gender', 'Role', 'Status', 'Type', 'Notes', 'Code', 'Number',
    'Username', 'Login', 'Access', 'Level', 'Environment', 'Service',
    'Region', 'Resource', 'Details', 'Summary', 'Version', 'Reference',
    'Signature', 'Designation', 'Location', 'Contact', 'Billing',
    'Development', 'Production', 'UAT', 'Staging',
}

# Date strings to exclude from password matching (Bug Fix 2)
# Any value matching a date pattern is never treated as a password
DATE_PATTERN = re.compile(
    r'^\d{1,2}[-\/\.]\w{2,9}[-\/\.]\d{2,4}$'   # 01-Jan-2025, 05/03/2025
    r'|\d{4}[-\/\.]\d{1,2}[-\/\.]\d{1,2}'        # 2025-01-05
    r'|\d{1,2}[-\/\.]\d{1,2}[-\/\.]\d{2,4}',     # 01-01-2025
    re.IGNORECASE
)

def is_safe_to_mask(val: str, pii_type: str) -> bool:
    """
    Returns False if the value should NOT be masked.
    Prevents false positives on column headers and date strings.
    """
    stripped = val.strip()

    # Bug Fix 1: never mask standalone column/header words
    if stripped in HEADER_WORDS:
        return False

    # Bug Fix 2: never treat date-formatted strings as passwords
    if pii_type == 'password' and DATE_PATTERN.match(stripped):
        return False

    # Never mask very short values (1-2 chars) unless they're specific types
    if len(stripped) <= 2 and pii_type not in ('in_aadhaar', 'uk_nino', 'us_ssn'):
        return False

    # Never mask placeholder/null values
    if stripped.upper() in ('N/A', 'TBD', 'NA', '-', '—', 'NONE', 'NULL', 'YES', 'NO'):
        return False

    return True


# ─── Regex detection (fast, always runs) ──────────────────────────────────────

# Bug Fix 6: High-priority patterns applied BEFORE short numeric patterns
# Order matters: longer/specific patterns must be checked before short ones
# (e.g. full phone number before ZIP, full API key before partial number)
HIGH_PRIORITY_TYPES = {
    'email', 'url', 'in_phone', 'au_phone', 'uk_phone', 'us_phone',
    'ca_phone', 'eu_phone', 'ae_phone', 'sa_phone', 'sg_phone',
    'api_key', 'aws_key', 'aws_secret', 'jwt_token',
    'in_aadhaar', 'in_gst', 'in_pan', 'uk_nino', 'uk_postcode',
    'us_ssn', 'ca_postcode', 'sg_nric', 'eu_iban', 'eu_vat',
    'credit_card', 'ip_address', 'ipv6',
}

def detect_pii_with_regex(text: str) -> list:
    """
    Global PII detection — IN, AU, UK, US, CA, EU, UAE, SG + universal.
    Bug Fix 6: high-priority (longer) patterns run first to prevent
    short patterns (ZIP, bank acct) from fragmenting longer strings.
    Bug Fix 1+2: safe-value guard prevents header words and dates being masked.
    """
    found = []

    # Pass 1: High-priority patterns first (phone, API keys, credentials)
    for pii_type in HIGH_PRIORITY_TYPES:
        pattern = REGEX_PATTERNS.get(pii_type)
        if not pattern:
            continue
        try:
            for match in re.finditer(pattern, text):
                val = match.group().strip()
                if len(val) > 1 and is_safe_to_mask(val, pii_type):
                    found.append({"text": val, "type": pii_type, "source": "Regex"})
        except re.error:
            continue

    # Collect already-matched spans to avoid double-matching
    matched_spans = set()
    for item in found:
        for m in re.finditer(re.escape(item["text"]), text):
            matched_spans.add((m.start(), m.end()))

    # Pass 2: Remaining patterns (postcodes, bank accts, IDs etc.)
    for pii_type, pattern in REGEX_PATTERNS.items():
        if pii_type in HIGH_PRIORITY_TYPES:
            continue
        try:
            for match in re.finditer(pattern, text):
                val = match.group().strip()
                if len(val) <= 1:
                    continue
                # Bug Fix 3 & 6: skip if this span overlaps a high-priority match
                span = (match.start(), match.end())
                overlaps = any(
                    not (span[1] <= s[0] or span[0] >= s[1])
                    for s in matched_spans
                )
                if overlaps:
                    continue
                if is_safe_to_mask(val, pii_type):
                    found.append({"text": val, "type": pii_type, "source": "Regex"})
                    matched_spans.add(span)
        except re.error:
            continue

    # Context-aware keyword: value extraction
    for ctx_pattern, (val_pattern, pii_type) in CONTEXT_LABELS.items():
        for ctx_match in re.finditer(ctx_pattern, text):
            rest = text[ctx_match.end():]
            val_match = re.match(val_pattern, rest.strip())
            if val_match:
                val = val_match.group().strip().rstrip(',;.')
                if is_safe_to_mask(val, pii_type):
                    found.append({"text": val, "type": pii_type, "source": "Regex"})

    # Proper name detection (global — honourifics + Title Case filtering)
    for pattern in NAME_PATTERNS:
        for match in re.finditer(pattern, text):
            words = match.group().strip().split()
            if any(w in NOT_NAMES for w in words):
                continue
            if len(words) < 2 or len(words) > 4:
                continue
            if all(w.isupper() for w in words):
                continue
            val = match.group().strip()
            if is_safe_to_mask(val, 'full_name'):
                found.append({"text": val, "type": "full_name", "source": "Regex"})

    return found


# ─── Core masking engine ───────────────────────────────────────────────────────

def build_replacement_map(full_text: str, model) -> tuple:
    """
    Core engine: regex (instant) + optional AI (max 3 chunks × 15s = 45s cap).
    Regex always returns. AI is a bonus — any crash/timeout is silently ignored.
    Bug Fix 6: replacement map sorted by length desc — longer strings replaced
    first so partial matches don't fragment longer PII values.
    """
    detections = []

    # STEP 1: Regex — always runs, instant
    for entity in detect_pii_with_regex(full_text):
        label = MASK_LABELS.get(entity["type"], f"[{entity['type'].upper()}]")
        detections.append({"original": entity["text"], "masked": label,
                           "type": entity["type"], "source": "Regex"})

    # STEP 2: AI — max 3 chunks of 500 chars, 15s timeout each
    if model:
        sample = full_text[:1500]
        chunks = [sample[i:i+LLM_CHUNK] for i in range(0, len(sample), LLM_CHUNK)]
        for chunk in chunks[:3]:
            if not chunk.strip():
                continue
            for entity in detect_pii_with_llm(chunk, model):
                label = MASK_LABELS.get(entity["type"], f"[{entity['type'].upper()}]")
                detections.append({"original": entity["text"], "masked": label,
                                   "type": entity["type"], "source": "AI"})

    # STEP 3: Build replacement map
    # Sort by length descending — replace longer strings FIRST (Bug Fix 6)
    # This ensures "ACb7c15432de8f..." is replaced before "123456" inside it
    replacement_map = {}
    seen = set()
    for entity in sorted(detections, key=lambda x: len(x["original"]), reverse=True):
        orig = entity["original"].strip()
        if not orig or orig in seen or len(orig) < 2:
            continue
        if orig.startswith('[') and orig.endswith(']'):
            continue
        # Final safe-value check before adding to map
        if not is_safe_to_mask(orig, entity["type"]):
            continue
        seen.add(orig)
        replacement_map[orig] = entity["masked"]

    return replacement_map, detections


def apply_replacements(text: str, replacement_map: dict) -> str:
    """
    Apply all replacements. Longer strings applied first (map already sorted).
    Bug Fix 6: sorted() here ensures correct order even if map order varies.
    """
    for orig in sorted(replacement_map.keys(), key=len, reverse=True):
        text = text.replace(orig, replacement_map[orig])
    return text


# ─── File Processors ───────────────────────────────────────────────────────────

def process_txt(file_path, model):
    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
        content = f.read()
    replacement_map, detections = build_replacement_map(content, model)
    masked = apply_replacements(content, replacement_map)
    out_path = file_path + "_masked.txt"
    with open(out_path, 'w', encoding='utf-8') as f:
        f.write(masked)
    return out_path, detections, masked[:500]


def process_docx(file_path, model):
    doc = Document(file_path)

    # Collect ALL text first — one AI call for entire doc
    full_text_parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text_parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        full_text_parts.append(para.text)

    full_text = "\n".join(full_text_parts)
    print(f"[Mask] Building replacement map from {len(full_text_parts)} paragraphs...")
    replacement_map, detections = build_replacement_map(full_text, model)
    print(f"[Mask] Found {len(replacement_map)} unique PII items — applying to document...")

    preview_parts = []
    for para in doc.paragraphs:
        if not para.text.strip():
            continue
        masked_text = apply_replacements(para.text, replacement_map)
        preview_parts.append(masked_text)
        if masked_text != para.text and para.runs:
            for run in para.runs:
                run.text = ""
            para.runs[0].text = masked_text

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        masked_text = apply_replacements(para.text, replacement_map)
                        if masked_text != para.text and para.runs:
                            for run in para.runs:
                                run.text = ""
                            para.runs[0].text = masked_text

    out_path = file_path + "_masked.docx"
    doc.save(out_path)
    return out_path, detections, "\n".join(preview_parts[:20])


def process_xlsx(file_path, model):
    wb = openpyxl.load_workbook(file_path)
    yellow_fill = PatternFill(start_color="FFFF00", end_color="FFFF00", fill_type="solid")

    all_text_parts = []
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if cell.value and isinstance(cell.value, str) and cell.value.strip():
                    all_text_parts.append(cell.value)

    full_text = "\n".join(all_text_parts)
    replacement_map, detections = build_replacement_map(full_text, model)

    preview_parts = []
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if cell.value and isinstance(cell.value, str) and cell.value.strip():
                    masked_text = apply_replacements(cell.value, replacement_map)
                    if masked_text != cell.value:
                        preview_parts.append(f"{cell.coordinate}: {cell.value} → {masked_text}")
                        cell.value = masked_text
                        cell.fill = yellow_fill

    out_path = file_path + "_masked.xlsx"
    wb.save(out_path)
    return out_path, detections, "\n".join(preview_parts[:20])


# ─── Routes ───────────────────────────────────────────────────────────────────

@app.route('/')
def index():
    return send_from_directory('static', 'index.html')


@app.route('/sanitize')
def sanitize():
    return send_from_directory('static', 'sanitize.html')


@app.route('/project-docs')
def project_docs():
    return send_from_directory('static', 'project_docs.html')


@app.route('/call-summarizer')
def call_summarizer():
    return send_from_directory('static', 'call_summarizer.html')


@app.route('/api/summarize-call', methods=['POST'])
@rate_limit("20 per minute; 200 per hour")
def summarize_call():
    data   = request.get_json() or {}
    source = data.get('source', '').strip()
    if not source:
        return jsonify({"error": "No transcript provided"}), 400

    model = OPENROUTER_MODEL if USE_CLOUD else get_model()
    if not model:
        return jsonify({"error": "No LLM backend. Set OPENROUTER_API_KEY or run: llama-server"}), 503

    # 2500 chars preserves names, owner assignments, and risk detail from transcript
    src_short = source[:2500].strip()
    prompt = """\
Role: Expert Business Analyst (15y exp).
Task: Analyze the call transcript below and return a structured JSON summary.
Output: ONLY valid JSON. No markdown, no prose.
Structure: {{"key_takeaways":["1-2 sentence insight"],"action_items":[{{"task":"What to do","owner":"Name or role from transcript (TBD if not mentioned)","deadline":"When (TBD if not mentioned)"}}],"decisions":["Decision made"],"risks":["Risk or blocker raised"],"next_steps":["Immediate follow-up"]}}
Constraints:
- key_takeaways: 10-15 items, each a distinct insight from the call
- action_items: extract all explicitly assigned tasks; use exact names from transcript for owner
- decisions: only firm decisions confirmed on the call
- risks: blockers, concerns, or dependencies raised — include specific details from transcript
- next_steps: concrete follow-ups in priority order, max 7
Every item must be traceable to the transcript. No invented content.
Transcript: {src}""".format(src=src_short)

    try:
        t   = _model_timeout(model)
        raw = _call_llm(model, prompt, timeout=t,
                        options_override={"max_tokens": 1200}, format_json=True)
        result = _parse_json_response(raw, 'summary')
        return jsonify({"success": True, "data": result, "model": model})
    except requests.exceptions.Timeout:
        return jsonify({"error": "Timed out. Try a shorter transcript or faster model."}), 504
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route('/api/status')
def status():
    backend_info = check_backend_available()
    active = OPENROUTER_MODEL if USE_CLOUD else get_model()
    return jsonify({
        "status": "ok",
        "backend": "openrouter" if USE_CLOUD else "llamacpp",
        "llama_server": backend_info,
        "active_model": active,
        "cloud_enabled": USE_CLOUD,
        "max_file_size_mb": 50,
        "ai_timeout_seconds": LLM_TIMEOUT,
        "chunk_size": LLM_CHUNK,
    })


@app.route('/api/models')
def list_models():
    """Return loaded llama-server model with profile hints for the dropdown."""
    try:
        r = requests.get(LLAMA_MODELS_URL, timeout=3)
        if r.status_code != 200:
            return jsonify({"models": [], "default": None, "error": "llama-server unreachable"}), 200
        raw = [m.get("id", "") for m in r.json().get("data", []) if m.get("id")]
        # Enrich with profile hints so UI can display time estimates
        enriched = []
        for name in raw:
            prof = _model_profile(name)
            enriched.append({
                "name":         name,
                "timeout":      prof["timeout"],
                "predict_mult": prof["predict_mult"],
            })
        default = get_model()
        return jsonify({"models": enriched, "default": default})
    except Exception as e:
        return jsonify({"models": [], "default": None, "error": str(e)}), 200


@app.route('/api/download-call-docx', methods=['POST'])
def download_call_docx():
    payload         = request.get_json() or {}
    d               = payload.get('data', {})
    transcript_name = payload.get('transcript_name', '').strip()

    safe_name = re.sub(r'[^\w\s-]', '', transcript_name) if transcript_name else ''
    filename  = f"Call Summary - {safe_name}.docx" if safe_name else "Call Summary.docx"

    doc = Document()
    title = doc.add_heading(f"Call Summary{' — ' + transcript_name if transcript_name else ''}", 0)

    sections = [
        ('key_takeaways',  '💡 Key Takeaways',    'list_numbered'),
        ('action_items',   '✅ Action Items',      'action'),
        ('decisions',      '🎯 Decisions Made',    'list_bullet'),
        ('risks',          '⚠️  Risks & Blockers', 'list_bullet'),
        ('next_steps',     '➡️  Next Steps',       'list_numbered'),
    ]

    for field, heading, style in sections:
        doc.add_heading(heading, level=1)
        items = d.get(field, [])
        if not items:
            doc.add_paragraph('None recorded.').italic = True
            continue
        for i, item in enumerate(items):
            if field == 'action_items':
                task     = item.get('task', str(item)) if isinstance(item, dict) else str(item)
                owner    = item.get('owner', 'TBD')    if isinstance(item, dict) else 'TBD'
                deadline = item.get('deadline', 'TBD') if isinstance(item, dict) else 'TBD'
                p = doc.add_paragraph()
                p.add_run(f"{i+1}. {task}").bold = True
                p.add_run(f"\n   Owner: {owner}  |  Deadline: {deadline}")
            elif style == 'list_numbered':
                doc.add_paragraph(f"{i+1}. {item}")
            else:
                doc.add_paragraph(f"• {item}")

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(tmp.name)
    tmp.close()
    return send_file(tmp.name, as_attachment=True, download_name=filename,
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


@app.route('/api/download-docx', methods=['POST'])
def download_docx():
    payload      = request.get_json() or {}
    data         = payload.get('data', payload)          # support both wrapped and flat
    project_name = payload.get('project_name', '').strip()
    artifact_label = payload.get('artifact_label', '').strip()

    # Build filename: "Risk Register - My Project.docx" or fallback
    safe_artifact = re.sub(r'[^\w\s-]', '', artifact_label or 'Agile Project Documents')
    safe_project  = re.sub(r'[^\w\s-]', '', project_name) if project_name else ''
    filename = f"{safe_artifact} - {safe_project}.docx" if safe_project else f"{safe_artifact}.docx"

    doc = Document()
    doc.add_heading(artifact_label or 'Agile Project Documents', 0)
    if project_name:
        doc.add_paragraph(f"Project: {project_name}").italic = True

    artifact_labels = {
        'backlog':       'Product Backlog',
        'sprint_plan':   'Sprint Planning',
        'sprint_review': 'Sprint Reviews',
        'retrospective': 'Retrospective',
        'risk_register': 'Risk Register',
        'test_cases':    'Test Cases',
    }

    for key, artifact_data in data.items():
        label = artifact_labels.get(key, key.replace('_', ' ').title())
        doc.add_heading(label, level=1)

        if key == 'backlog':
            # Handle list (after _extract_complete_objects), dict, or _raw fallback
            if isinstance(artifact_data, list):
                epics = artifact_data
            elif isinstance(artifact_data, dict) and '_raw' in artifact_data:
                try:
                    p = json.loads(artifact_data['_raw'])
                    epics = p if isinstance(p, list) else p.get('epics', [])
                except Exception:
                    epics = []
            else:
                epics = artifact_data.get('epics', []) if isinstance(artifact_data, dict) else []
            for epic in epics:
                doc.add_heading(f"{epic.get('id','')} {epic.get('title','Epic')}".strip(), level=2)
                if epic.get('description'):
                    doc.add_paragraph(epic['description'])
                for s in epic.get('stories', []):
                    pts = s.get('story_points', s.get('points', '?'))
                    pri = s.get('priority', '')
                    p = doc.add_paragraph()
                    p.add_run(f"{s.get('id','')} {s.get('title','')}").bold = True
                    p.add_run(f"  [{pts} SP]  [{pri}]")
                    if s.get('description'):
                        doc.add_paragraph(s['description'], style='List Bullet')
                    for ac in s.get('acceptance_criteria', []):
                        doc.add_paragraph(f"\u2713 {ac}", style='List Bullet')

        elif key == 'sprint_plan':
            if isinstance(artifact_data, list):
                sprints = artifact_data
            elif isinstance(artifact_data, dict) and '_raw' in artifact_data:
                try:
                    p = json.loads(artifact_data['_raw'])
                    sprints = p if isinstance(p, list) else p.get('sprints', [])
                except Exception:
                    sprints = []
            else:
                sprints = artifact_data.get('sprints', []) if isinstance(artifact_data, dict) else []
            for s in sprints:
                doc.add_heading(f"Sprint {s.get('number','')}".strip(), level=2)
                if s.get('goal'):
                    doc.add_paragraph(f"Goal: {s['goal']}")
                doc.add_paragraph(f"Story Points: {s.get('total_points', s.get('points', 0))}")
                stories = s.get('stories', [])
                if stories:
                    ids = [st if isinstance(st, str) else st.get('id', str(st)) for st in stories]
                    doc.add_paragraph("Stories: " + ", ".join(ids))
                if s.get('deliverable'):
                    doc.add_paragraph(f"Deliverable: {s['deliverable']}")

        elif key == 'sprint_review':
            reviews = artifact_data if isinstance(artifact_data, list) else artifact_data.get('reviews', [])
            for r in reviews:
                doc.add_heading(f"Sprint {r.get('sprint','')} Review".strip(), level=2)
                if r.get('planned_stories'):
                    doc.add_paragraph("Planned: " + ", ".join(str(s) for s in r['planned_stories']))
                if r.get('completed_stories'):
                    doc.add_paragraph("Completed: " + ", ".join(str(s) for s in r['completed_stories']))
                if r.get('demo_notes'):
                    doc.add_paragraph(f"Demo Notes: {r['demo_notes']}")
                if r.get('feedback_template'):
                    doc.add_paragraph(f"Stakeholder Feedback: {r['feedback_template']}")
                if r.get('next_sprint_adjustments'):
                    doc.add_paragraph(f"Next Sprint Adjustments: {r['next_sprint_adjustments']}")

        elif key == 'retrospective':
            for field, section_label in [('went_well','What Went Well'),('improve','What to Improve'),
                                          ('action_items','Action Items'),('team_health_check','Team Health Check')]:
                items = artifact_data.get(field, []) if isinstance(artifact_data, dict) else []
                if items:
                    doc.add_heading(section_label, level=2)
                    for item in items:
                        doc.add_paragraph(f"\u2022 {item}")

        elif key == 'risk_register':
            # Recover if server returned _raw (JSON parse failed at generation time)
            if isinstance(artifact_data, dict) and '_raw' in artifact_data:
                try:
                    parsed = json.loads(artifact_data['_raw'])
                    artifact_data = parsed if isinstance(parsed, list) else parsed.get('risks', parsed)
                except Exception:
                    pass
            risks = artifact_data if isinstance(artifact_data, list) else artifact_data.get('risks', [])
            for r in risks:
                title = f"{r.get('id','')} {r.get('title', r.get('risk','Risk'))}".strip()
                doc.add_heading(title, level=2)
                if r.get('description'):
                    doc.add_paragraph(r['description'])
                doc.add_paragraph(
                    f"Probability: {r.get('probability','M')}  |  Impact: {r.get('impact','M')}  |  Severity: {r.get('severity','Medium')}")
                if r.get('mitigation'):
                    doc.add_paragraph(f"Mitigation: {r['mitigation']}")

        elif key == 'test_cases':
            cases = artifact_data if isinstance(artifact_data, list) else artifact_data.get('test_cases', [])
            for tc in cases:
                doc.add_heading(f"{tc.get('id','')} {tc.get('title','Test Case')}".strip(), level=2)
                doc.add_paragraph(f"Type: {tc.get('type','')}  |  Story: {tc.get('story_id','')}  |  Priority: {tc.get('priority','')}")
                if tc.get('precondition'):
                    doc.add_paragraph(f"Precondition: {tc['precondition']}")
                for i, step in enumerate(tc.get('steps', []), 1):
                    doc.add_paragraph(f"{i}. {step}", style='List Number')
                if tc.get('expected_result'):
                    doc.add_paragraph(f"Expected Result: {tc['expected_result']}")

        else:
            txt = artifact_data if isinstance(artifact_data, str) else json.dumps(artifact_data, indent=2)
            doc.add_paragraph(txt)

        doc.add_page_break()

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(tmp.name)
    tmp.close()
    return send_file(
        tmp.name,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


@app.route('/api/mask', methods=['POST'])
def mask_file():
    if 'file' not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files['file']
    if not file.filename:
        return jsonify({"error": "Empty filename"}), 400

    file.seek(0, 2)
    size = file.tell()
    file.seek(0)
    if size > MAX_FILE_SIZE:
        return jsonify({"error": "File too large. Max 50MB allowed."}), 400

    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in ['.txt', '.docx', '.xlsx', '.xls', '.csv']:
        return jsonify({"error": f"Unsupported format: {ext}. Supported: txt, docx, xlsx, csv"}), 400

    # UUID prefix prevents filename collision when concurrent users upload same filename
    safe_ext  = os.path.splitext(file.filename)[1].lower()
    tmp_name  = f"{uuid.uuid4().hex}{safe_ext}"
    tmp_path  = os.path.join(UPLOAD_FOLDER, tmp_name)
    file.save(tmp_path)

    model = get_model()
    print(f"[Mask] File={file.filename} | AI model={'None (regex only)' if not model else model}")

    try:
        if ext in ['.txt', '.csv']:
            out_path, detections, preview = process_txt(tmp_path, model)
        elif ext == '.docx':
            out_path, detections, preview = process_docx(tmp_path, model)
        elif ext in ['.xlsx', '.xls']:
            out_path, detections, preview = process_xlsx(tmp_path, model)

        seen_orig = set()
        unique_detections = []
        for d in detections:
            if d["original"] not in seen_orig:
                seen_orig.add(d["original"])
                unique_detections.append(d)

        return jsonify({
            "success": True,
            "original_filename": file.filename,
            "masked_filename": os.path.basename(out_path),
            "detections": unique_detections,
            "detection_count": len(unique_detections),
            "preview": preview[:500],
            "download_token": os.path.basename(out_path),
            "ai_used": model is not None,
            "model": model or "regex-only"
        })

    except Exception as e:
        print(f"[Error] {e}")
        return jsonify({"error": str(e)}), 500


@app.route('/api/download/<filename>')
def download_file(filename):
    safe_name = os.path.basename(filename)
    file_path = os.path.join(UPLOAD_FOLDER, safe_name)
    if not os.path.exists(file_path):
        return jsonify({"error": "File not found"}), 404
    response = send_file(file_path, as_attachment=True, download_name=safe_name)
    # Clean up after serving — prevent disk accumulation on long-running pod
    try:
        os.unlink(file_path)
    except Exception:
        pass
    return response


@app.route('/api/extract-text', methods=['POST'])
@rate_limit("20 per minute; 300 per hour")
def extract_text_api():
    """Server-side text extraction for binary formats (docx, pdf, xlsx).
    Browser cannot read these as plain text — readAsText() returns raw ZIP bytes."""
    if 'file' not in request.files:
        return jsonify({"error": "No file provided"}), 400
    f = request.files['file']
    name = (f.filename or '').lower()

    try:
        if name.endswith('.docx'):
            from docx import Document as DocxDocument
            import io
            doc = DocxDocument(io.BytesIO(f.read()))
            parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text.strip())
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        parts.append(' | '.join(cells))
            text = '\n'.join(parts)
            if not text.strip():
                return jsonify({"error": "Document appears empty or unreadable"}), 400
            return jsonify({"text": text})

        elif name.endswith('.pdf'):
            data = f.read()
            # Try pdfplumber first, then pypdf
            try:
                import pdfplumber, io
                with pdfplumber.open(io.BytesIO(data)) as pdf:
                    text = '\n'.join(page.extract_text() or '' for page in pdf.pages)
                return jsonify({"text": text})
            except ImportError:
                pass
            try:
                import pypdf, io
                reader = pypdf.PdfReader(io.BytesIO(data))
                text = '\n'.join(page.extract_text() or '' for page in reader.pages)
                return jsonify({"text": text})
            except ImportError:
                return jsonify({"error": "PDF extraction unavailable. Install: pip install pdfplumber"}), 500

        elif name.endswith(('.xlsx', '.xls')):
            try:
                import openpyxl, io
                wb = openpyxl.load_workbook(io.BytesIO(f.read()), data_only=True)
                parts = []
                for sheet in wb.worksheets:
                    parts.append(f'[Sheet: {sheet.title}]')
                    for row in sheet.iter_rows(values_only=True):
                        cells = [str(c) for c in row if c is not None]
                        if cells:
                            parts.append(' | '.join(cells))
                return jsonify({"text": '\n'.join(parts)})
            except ImportError:
                return jsonify({"error": "Excel extraction unavailable. Install: pip install openpyxl"}), 500

        else:
            # txt, csv, md — plain text
            text = f.read().decode('utf-8', errors='replace')
            return jsonify({"text": text})

    except Exception as e:
        print(f"[extract-text] error: {e}")
        return jsonify({"error": str(e)}), 500


@app.route('/api/mask-text', methods=['POST'])
def mask_text_api():
    data = request.json
    if not data or 'text' not in data:
        return jsonify({"error": "No text provided"}), 400
    text = data['text'][:10000]
    model = get_model()
    replacement_map, detections = build_replacement_map(text, model)
    masked = apply_replacements(text, replacement_map)
    return jsonify({
        "masked_text": masked,
        "detections": detections,
        "detection_count": len(replacement_map)
    })


# ── Per-artifact focused prompts (small = fast) ────────────────────────────────

_ARTIFACT_PROMPTS = {
    'backlog': """\
Role: Senior Scrum Master (15y exp).
Task: Generate JSON Product Backlog for the project in source.
Output: ONLY valid JSON. No markdown, no prose.
Structure: {{"epics":[{{"id":"EP-01","title":"Epic title","description":"What this epic covers","stories":[{{"id":"US-01-01","title":"As a [role] I want [goal] so that [benefit]","story_points":3,"priority":"Must Have","acceptance_criteria":["Criterion 1","Criterion 2","Criterion 3"]}}]}}]}}
Constraints: 5 epics. 4 stories per epic. Points: 1/2/3/5/8/13. Priority: Must Have/Should Have/Could Have/Won't Have. 2-3 AC per story.
Focus: Derive every epic and story directly from features, integrations, user roles, and performance targets in source. No generic items.
Project info: {src}""",

    'sprint_plan': """\
Role: Senior Scrum Master (15y exp).
Task: Generate JSON Sprint Plan for the project in source. Velocity 28 pts, 2-week sprints.
Output: ONLY valid JSON. No markdown, no prose.
Structure: {{"duration_weeks":2,"team_velocity":28,"sprints":[{{"number":1,"goal":"Specific sprint goal tied to source deliverables","stories":["US-01-01","US-01-02","US-02-01"],"total_points":27,"deliverable":"What is demo-ready at end of sprint","risks":"Top risk for this sprint"}}]}}
Constraints: 4-5 sprints covering all stories. Each sprint: specific goal, 5-6 story IDs, points near velocity, one demo deliverable, one risk. Reference actual source deliverables.
Project info: {src}""",

    'sprint_review': """\
Role: Senior Scrum Master (15y exp).
Task: Generate JSON Sprint Review templates for the project in source.
Output: ONLY valid JSON. No markdown, no prose.
Structure: {{"reviews":[{{"sprint":1,"planned_stories":["US-01-01","US-01-02"],"completed_stories":["US-01-01"],"incomplete_stories":["US-01-02"],"demo_notes":"What to demo and how, naming actual features","stakeholder_feedback":"Specific questions for stakeholders named in source","next_sprint_adjustments":"Concrete changes based on this sprint outcome","velocity_actual":24,"velocity_planned":28}}]}}
Constraints: One review per sprint (3-4 reviews). Each: planned vs completed, specific demo notes naming features, stakeholder questions referencing source, concrete next-sprint adjustments.
Project info: {src}""",

    'retrospective': """\
Role: Senior Scrum Master (15y exp).
Task: Generate JSON Start-Stop-Continue Retrospective for the project in source.
Output: ONLY valid JSON. No markdown, no prose.
Structure: {{"went_well":["Item"],"improve":["Item"],"action_items":[{{"action":"Specific action","owner":"Role","due":"Sprint N"}}],"team_health_check":["Dimension: rating and note"],"process_improvements":["Improvement"]}}
Constraints: went_well: 5 specific items. improve: 5 specific items. action_items: 4-5 with owner and due sprint. team_health_check: 4 dimensions with rating. process_improvements: 3-4 items.
Focus: Reference actual technical challenges, integration risks, stakeholder dynamics, and team constraints from source. No generic Agile platitudes.
Project info: {src}""",

    'risk_register': """\
Role: Senior Scrum Master (15y exp).
Task: Generate JSON Risk Register for the project in source.
Output: ONLY valid JSON. No markdown, no prose.
Structure: {{"risks":[{{"id":"R-01","category":"Scope Creep","title":"Brief risk title","description":"One specific sentence from source.","probability":"High","impact":"High","severity":"Critical","mitigation":"One concrete sentence.","owner":"Role","sprint_impact":"Sprint 1-2"}}]}}
Constraints: 6 risks. Cover categories: Scope Creep, Resource Availability, Technical Integration, Security/Compliance, Timeline/Delivery, Team Capacity. severity: Critical/High/Medium/Low. probability/impact: High/Medium/Low. Values from source only.
Project info: {src}""",

    'test_cases': """\
Role: Senior QA Lead (15y exp).
Task: Generate JSON Test Suite for the project in source.
Output: ONLY valid JSON. No markdown, no prose.
Structure: {{"test_cases":[{{"id":"TC-01","title":"Test title","type":"Unit","feature":"Feature under test","story_id":"US-01-01","steps":["Step 1","Step 2","Step 3","Step 4"],"expected_result":"Specific expected outcome","priority":"High"}}]}}
Constraints: 10 test cases. type: Unit/Integration/UAT (at least 3 of each). 3-4 steps per test. Cover core features, API integrations, edge cases, and security requirements from source.
Project info: {src}""",
}


def _model_timeout(model_name):
    """Return appropriate request timeout in seconds based on model size.
    Handles both full model names (qwen2.5:7b-instruct-q4_K_M) and
    short aliases (qwen2.5) set via --alias at llama-server startup.
    """
    name = (model_name or "").lower()
    if "gemma4" in name:
        return 360
    if any(x in name for x in ("27b", "13b", "12b", "9b")):
        return 600
    if "7b" in name or "8b" in name:
        return 480
    # Known 7B aliases without size suffix
    if any(x in name for x in ("qwen2.5", "qwen3", "llama3", "mistral")):
        return 480
    if "4b" in name:
        return 240
    return 120   # tiny models (1b, tinyllama)


def _model_profile(model_name):
    """
    Per-model tuning profile. Tokens/sec and JSON reliability differ by model,
    so max_tokens + timeout scale accordingly. Returned dict applied on top of
    base _ARTIFACT_PREDICT by caller.
    """
    name = (model_name or "").lower()
    # Defaults: qwen2.5:7b sizing
    profile = {
        "predict_mult": 1.0,
        "timeout":      _model_timeout(model_name),
        "top_k":        1,
        "temperature":  0,
    }
    if "gemma4" in name:
        profile["top_k"]       = 40
        profile["temperature"] = 0.1
    if "gemma3" in name and "4b" in name:
        # Faster (~13 tok/s), smaller ctx sweet spot — shrink predict 20%
        profile["predict_mult"] = 0.85
    if "1b" in name or "tinyllama" in name:
        profile["predict_mult"] = 0.6
    if "27b" in name or "13b" in name or "12b" in name:
        profile["predict_mult"] = 1.1
    return profile


def _call_openrouter(prompt, timeout=60, format_json=False, max_tokens=1500):
    """
    OpenRouter cloud inference — free tier Gemma, OpenAI-compatible.
    Sub-10s per artifact vs. 2-3 min local.
    """
    headers = {
        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
        "Content-Type":  "application/json",
        "HTTP-Referer":  "http://localhost:5000",   # required by OpenRouter
        "X-Title":       "Agile Suite — Project Docs",
    }
    payload = {
        "model":       OPENROUTER_MODEL,
        "messages":    [{"role": "user", "content": prompt}],
        "temperature": 0.1,
        "max_tokens":  max_tokens,
    }
    if format_json:
        payload["response_format"] = {"type": "json_object"}
    r = requests.post(OPENROUTER_URL, headers=headers, json=payload, timeout=timeout)
    r.raise_for_status()
    j = r.json()
    return j["choices"][0]["message"]["content"]


def _call_llm(model, prompt, timeout=None, options_override=None, format_json=False):
    """Dispatch to OpenRouter if API key set, else llama-server local.
    Semaphore-guarded: caps global concurrent LLM calls to LLM_MAX_CONCURRENT
    matching llama-server --parallel slots to prevent VRAM thrash.
    """
    with _LLM_SEM:
        if USE_CLOUD:
            max_tok = (options_override or {}).get("max_tokens", 1500)
            t       = timeout or 60
            try:
                return _call_openrouter(prompt, timeout=t, format_json=format_json, max_tokens=max_tok)
            except Exception as e:
                print(f"[openrouter] failed ({e}), falling back to llama-server")
        return _call_llama_server(model, prompt, timeout=timeout,
                                  options_override=options_override, format_json=format_json)


def _stream_llama_server(messages, model, timeout=120):
    """Generator: yields text tokens from llama-server SSE stream (multi-turn messages)."""
    payload = {
        "model":       model,
        "messages":    messages,
        "max_tokens":  2048,
        "temperature": 0.7,
        "stream":      True,
    }
    with _LLAMA_SESSION.post(LLAMA_CHAT_URL, json=payload, stream=True, timeout=timeout) as r:
        r.raise_for_status()
        for line in r.iter_lines():
            if not line:
                continue
            decoded = line.decode("utf-8") if isinstance(line, bytes) else line
            if not decoded.startswith("data: "):
                continue
            data_str = decoded[6:].strip()
            if data_str == "[DONE]":
                return
            try:
                chunk = json.loads(data_str)
                token = chunk["choices"][0]["delta"].get("content", "")
                if token:
                    yield token
            except (KeyError, json.JSONDecodeError):
                pass


def _stream_openrouter(messages, timeout=120):
    """Generator: yields text tokens from OpenRouter SSE stream (multi-turn messages)."""
    headers = {
        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
        "Content-Type":  "application/json",
        "HTTP-Referer":  "http://localhost:5000",
        "X-Title":       "Agile Suite — Chat",
    }
    payload = {
        "model":       OPENROUTER_MODEL,
        "messages":    messages,
        "max_tokens":  2048,
        "temperature": 0.7,
        "stream":      True,
    }
    with requests.post(OPENROUTER_URL, headers=headers, json=payload, stream=True, timeout=timeout) as r:
        r.raise_for_status()
        for line in r.iter_lines():
            if not line:
                continue
            decoded = line.decode("utf-8") if isinstance(line, bytes) else line
            if not decoded.startswith("data: "):
                continue
            data_str = decoded[6:].strip()
            if data_str == "[DONE]":
                return
            try:
                chunk = json.loads(data_str)
                token = chunk["choices"][0]["delta"].get("content", "")
                if token:
                    yield token
            except (KeyError, json.JSONDecodeError):
                pass


def _call_llama_server(model, prompt, timeout=None, options_override=None, format_json=False):
    """
    Call llama-server via OpenAI-compatible /v1/chat/completions.

    Inference tuning (context size, GPU layers, threads) are llama-server
    startup flags — not per-request. Only generation params sent here:
      temperature=0     — deterministic, no sampling overhead
      top_k=1           — greedy decode, fastest sampling
      top_p=1.0         — disable nucleus sampling overhead
      repeat_penalty=1.0 — no repetition check overhead

    Pass options_override dict to tune max_tokens per task.
    """
    if timeout is None:
        timeout = _model_timeout(model)

    params = {
        "max_tokens":     1200,   # default ceiling; overridden per-artifact
        "temperature":    0,      # deterministic
        "top_k":          1,      # greedy decode
        "top_p":          1.0,    # disable nucleus sampling
        "repeat_penalty": 1.0,    # disable repetition penalty
    }
    if options_override:
        params.update(options_override)

    payload = {
        "model":    model,
        "messages": [{"role": "user", "content": prompt}],
        **params,
    }
    if format_json:
        payload["response_format"] = {"type": "json_object"}

    r = requests.post(LLAMA_CHAT_URL, json=payload, timeout=timeout)
    r.raise_for_status()
    return r.json()["choices"][0]["message"]["content"]


def _close_truncated_json(text):
    """Add missing closing brackets/braces to truncated JSON string.
    Handles mid-string truncation by stripping the incomplete string token."""
    stack = []
    in_string = False
    escape = False
    last_closed_pos = 0  # position after last fully-balanced close (stack empty)
    string_start = -1   # position of opening " of the current string token

    for i, ch in enumerate(text):
        if escape:
            escape = False
            continue
        if ch == '\\' and in_string:
            escape = True
            continue
        if ch == '"':
            if in_string:
                in_string = False
                string_start = -1
                if not stack:
                    last_closed_pos = i + 1
            else:
                in_string = True
                string_start = i
            continue
        if in_string:
            continue
        if ch in ('{', '['):
            stack.append('}' if ch == '{' else ']')
        elif ch in ('}', ']'):
            if stack and stack[-1] == ch:
                stack.pop()
                if not stack:
                    last_closed_pos = i + 1

    if in_string:
        # Truncated mid-string — strip the incomplete string token and recurse
        if string_start > 0:
            # Remove incomplete string + trailing comma/whitespace before it
            prefix = text[:string_start].rstrip().rstrip(',').rstrip()
            # If prefix now ends with `:` (incomplete key-value pair), strip back
            # past the key too — walk back until we're past the key string + comma/opening brace
            if prefix.endswith(':'):
                # Strip: " : → find preceding " (end of key), then preceding " (start of key)
                end_quote = prefix.rfind('"', 0, len(prefix) - 1)
                if end_quote > 0:
                    start_quote = prefix.rfind('"', 0, end_quote)
                    if start_quote > 0:
                        prefix = prefix[:start_quote].rstrip().rstrip(',').rstrip()
            if prefix:
                return _close_truncated_json(prefix)
        # Fall back to last fully-closed position if available
        if last_closed_pos > 0:
            return _close_truncated_json(text[:last_closed_pos])
        return text  # cannot repair
    if not stack:
        return text  # already balanced
    return text + ''.join(reversed(stack))


def _extract_complete_objects(text, array_key):
    """Last-resort: scan text for complete {...} objects inside a named array.
    Works even when outer structure is truncated. Returns {array_key: [obj,...]} or None."""
    import re
    m = re.search(rf'"{re.escape(array_key)}"\s*:\s*\[', text)
    if not m:
        return None
    pos = m.end()
    objects = []
    depth = 0
    in_string = False
    escape = False
    obj_start = None

    for i in range(pos, len(text)):
        ch = text[i]
        if escape:
            escape = False
            continue
        if ch == '\\' and in_string:
            escape = True
            continue
        if ch == '"':
            in_string = not in_string
            continue
        if in_string:
            continue
        if ch == '{':
            if depth == 0:
                obj_start = i
            depth += 1
        elif ch == '}':
            depth -= 1
            if depth == 0 and obj_start is not None:
                try:
                    obj = json.loads(text[obj_start:i + 1])
                    objects.append(obj)
                except json.JSONDecodeError:
                    pass
                obj_start = None
        elif ch == ']' and depth == 0:
            break  # end of array

    if objects:
        print(f"[parse] extracted {len(objects)} complete objects from '{array_key}' array")
        return {array_key: objects}
    return None


# Keys used in each artifact's JSON output — for targeted object extraction
_ARTIFACT_ARRAY_KEYS = {
    'backlog':        'epics',
    'sprint_plan':    'sprints',
    'sprint_review':  'reviews',
    'retrospective':  None,          # flat structure, no array key needed
    'risk_register':  'risks',
    'test_cases':     'test_cases',
}


def _parse_json_response(raw, artifact_key):
    """Best-effort JSON extraction with truncation repair. Returns a dict/list or _raw fallback."""
    print(f"[parse] {artifact_key}: raw_len={len(raw)} first80={repr(raw[:80])}")
    cleaned = raw.strip()

    # ── Strip markdown fences ──────────────────────────────────────────────────
    if cleaned.startswith("```"):
        lines = cleaned.split("\n")
        cleaned = "\n".join(lines[1:])
        if cleaned.lstrip().startswith("json"):
            cleaned = cleaned.lstrip()[4:]
    if cleaned.endswith("```"):
        cleaned = "\n".join(cleaned.split("\n")[:-1])
    cleaned = cleaned.strip()

    # ── Unwrap single-key wrappers e.g. {"risks":[...]} → list ────────────────
    def _unwrap(obj):
        if not isinstance(obj, dict):
            return obj
        for k in (artifact_key, 'risks', 'epics', 'sprints', 'reviews',
                  'test_cases', 'items', 'data', 'backlog'):
            if k in obj and len(obj) == 1:
                return obj[k]
        return obj

    def _try(fragment):
        try:
            return _unwrap(json.loads(fragment))
        except json.JSONDecodeError:
            return None

    # ── Pass 1: direct parse of outermost { } or [ ] ─────────────────────────
    for start_ch, end_ch in [('{', '}'), ('[', ']')]:
        s = cleaned.find(start_ch)
        e = cleaned.rfind(end_ch)
        if s != -1 and e != -1 and e > s:
            result = _try(cleaned[s:e + 1])
            if result is not None:
                print(f"[parse] {artifact_key}: direct parse OK")
                return result

    # ── Pass 2: close truncated JSON then parse ───────────────────────────────
    for start_ch in ('{', '['):
        s = cleaned.find(start_ch)
        if s != -1:
            repaired = _close_truncated_json(cleaned[s:])
            result = _try(repaired)
            if result is not None:
                print(f"[parse] {artifact_key}: truncation-close repaired OK")
                return result

    # ── Pass 3: extract complete objects from named array (handles deep truncation) ─
    array_key = _ARTIFACT_ARRAY_KEYS.get(artifact_key)
    if array_key:
        result = _extract_complete_objects(cleaned, array_key)
        if result is not None:
            return _unwrap(result)

    print(f"[parse] {artifact_key}: all passes failed → _raw")
    return {"_raw": cleaned}


# ── /api/refine-artifact  (re-generate with feedback context) ───────────────────
@app.route('/api/refine-artifact', methods=['POST'])
def refine_artifact():
    """Re-generate an artifact using the existing output as context for improvement."""
    data = request.json or {}
    artifact  = data.get('artifact', '')
    source    = data.get('source', '')
    current   = data.get('current_output', '')   # JSON string of current output
    feedback  = data.get('feedback', '')          # optional user feedback text

    if not artifact or not source:
        return jsonify({"error": "artifact and source required"}), 400
    if artifact not in _ARTIFACT_PROMPTS:
        return jsonify({"error": f"Unknown artifact: {artifact}"}), 400

    model = OPENROUTER_MODEL if USE_CLOUD else get_model()
    if not model:
        return jsonify({"error": "No LLM backend. Set OPENROUTER_API_KEY or start llama-server"}), 503

    _ARTIFACT_PREDICT = {
        'backlog': 1500, 'sprint_plan': 900, 'sprint_review': 1000,
        'retrospective': 800, 'risk_register': 1100, 'test_cases': 1300,
    }
    predict  = _ARTIFACT_PREDICT.get(artifact, 1000)
    src_short = source[:3000].strip()

    artifact_labels = {
        'backlog': 'Product Backlog', 'sprint_plan': 'Sprint Planning',
        'sprint_review': 'Sprint Reviews', 'retrospective': 'Retrospective',
        'risk_register': 'Risk Register', 'test_cases': 'Test Cases',
    }
    label = artifact_labels.get(artifact, artifact.replace('_', ' ').title())

    # Summarise current output (first 800 chars to stay within context)
    current_summary = current[:800].strip() if current else ''
    feedback_line   = f'\nUser feedback: {feedback.strip()}' if feedback and feedback.strip() else ''

    refine_prompt = (
        f"Review the previously generated {label} shown below and produce a MORE DETAILED, "
        f"COMPREHENSIVE version with additional relevant entries specific to the project."
        f"{feedback_line}\n"
        f"Previous output (partial): {current_summary}\n\n"
        f"Now generate a complete, improved {label}.\n\n"
        + _ARTIFACT_PROMPTS[artifact].format(src=src_short)
    )

    try:
        t   = _model_timeout(model)
        print(f"[refine] key={artifact} model={model} predict={predict} timeout={t}s")
        raw    = _call_llm(model, refine_prompt, timeout=t,
                           options_override={"max_tokens": predict}, format_json=True)
        result = _parse_json_response(raw, artifact)
        return jsonify({"success": True, "artifact": artifact, "data": result, "model": model})
    except requests.exceptions.Timeout:
        return jsonify({"error": "Timed out. Try a faster model."}), 504
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ── /api/generate-artifact  (generates ONE artifact at a time) ──────────────────
@app.route('/api/generate-artifact', methods=['POST'])
@rate_limit("30 per minute; 300 per hour")
def generate_artifact():
    """Generate a single Agile artifact. Called N times by the frontend."""
    data = request.json or {}
    source       = data.get('source', '')
    artifact     = data.get('artifact', '')
    model_override = (data.get('model') or '').strip()

    if not source or not artifact:
        return jsonify({"error": "source and artifact are required"}), 400
    if artifact not in _ARTIFACT_PROMPTS:
        return jsonify({"error": f"Unknown artifact: {artifact}"}), 400

    # Priority: UI override > cloud > auto-pick
    if model_override:
        model = model_override
    else:
        model = OPENROUTER_MODEL if USE_CLOUD else get_model()
    if not model:
        return jsonify({
            "error": "No LLM backend available. Set OPENROUTER_API_KEY env var OR start llama-server"
        }), 503

    # Per-artifact max_tokens — sized for qwen2.5:7b at ~10 tok/s on M4
    # Richer outputs since qwen2.5 has better JSON adherence
    # llama-server --parallel 4 → 6 artifacts in ~220s total
    _ARTIFACT_PREDICT = {
        'backlog':        1500,  # 5E × 4S + AC ≈ 1400 tok ≈ 140s
        'sprint_plan':     900,  # 4-5 sprints ≈ 850 tok ≈ 85s
        'sprint_review':  1000,  # 4 reviews ≈ 950 tok ≈ 95s
        'retrospective':   800,  # 4 sections × 5 items ≈ 750 tok ≈ 75s
        'risk_register':  1100,  # 8 risks × 7 fields ≈ 1050 tok ≈ 105s
        'test_cases':     1300,  # 10 cases × 4 steps ≈ 1250 tok ≈ 125s
    }
    base_predict = _ARTIFACT_PREDICT.get(artifact, 1000)
    profile      = _model_profile(model)
    predict      = max(400, int(base_predict * profile["predict_mult"]))

    # 8000 chars ≈ 2000 tokens input; leaves ~14K tokens for output within num_ctx=16384
    # Gemma 4 handles 128K ctx natively — we could go much higher but 8K source is usually sufficient
    src_short = source[:3000].strip()
    prompt    = _ARTIFACT_PROMPTS[artifact].format(src=src_short)

    try:
        t = profile["timeout"]
        print(f"[artifact] key={artifact} model={model} src_len={len(src_short)} max_tokens={predict} timeout={t}s")
        raw    = _call_llm(model, prompt, timeout=t, options_override={"max_tokens": predict}, format_json=True)
        result = _parse_json_response(raw, artifact)
        print(f"[artifact] key={artifact} done, result_type={type(result).__name__}")
        return jsonify({"success": True, "artifact": artifact, "data": result, "model": model})
    except requests.exceptions.Timeout:
        return jsonify({"error": f"Timed out generating {artifact}. Try a faster model."}), 504
    except Exception as e:
        print(f"[artifact error] {artifact}: {e}")
        return jsonify({"error": str(e)}), 500


# ── /api/generate-docs  (kept for backwards compat — calls generate-artifact N×) ─
@app.route('/api/generate-docs', methods=['POST'])
@rate_limit("5 per minute; 60 per hour")
def generate_docs():
    data = request.json or {}
    source    = data.get('source', '')
    artifacts = data.get('artifacts', list(_ARTIFACT_PROMPTS.keys()))

    model = OPENROUTER_MODEL if USE_CLOUD else get_model()
    if not model:
        return jsonify({"error": "No LLM backend. Set OPENROUTER_API_KEY or start llama-server"}), 503

    documents = {}
    for key in artifacts:
        if key not in _ARTIFACT_PROMPTS:
            continue
        src_short = source[:1500].strip()
        prompt    = _ARTIFACT_PROMPTS[key].format(src=src_short)
        try:
            raw  = _call_llm(model, prompt, timeout=90)
            documents[key] = _parse_json_response(raw, key)
        except Exception as e:
            print(f"[generate-docs] {key} failed: {e}")
            documents[key] = {"_raw": str(e)}

    return jsonify({"success": True, "documents": documents, "model": model})


# ── Chat routes ────────────────────────────────────────────────────────────────

@app.route('/chat')
def chat_page():
    return send_from_directory('static', 'chat.html')


@app.route('/api/chat/stream', methods=['POST'])
@rate_limit("20 per minute; 200 per hour")
def chat_stream():
    data       = request.json or {}
    session_id = data.get('session_id', '').strip()
    message    = data.get('message', '').strip()

    if not message:
        return jsonify({"error": "message required"}), 400
    if not session_id:
        return jsonify({"error": "session_id required"}), 400

    if session_id not in _CHAT_SESSIONS:
        _CHAT_SESSIONS[session_id] = []

    _CHAT_SESSIONS[session_id].append({"role": "user", "content": message})
    messages = [{"role": "system", "content": _CHAT_SYSTEM_PROMPT}] + _CHAT_SESSIONS[session_id]
    model    = OPENROUTER_MODEL if USE_CLOUD else (get_model() or LLM_MODEL)

    def generate():
        full_response = []
        try:
            with _LLM_SEM:
                token_gen = _stream_openrouter(messages) if USE_CLOUD else _stream_llama_server(messages, model)
                for token in token_gen:
                    full_response.append(token)
                    yield f"data: {json.dumps({'token': token})}\n\n"
            # Normal completion — save assistant response, then signal done
            if full_response:
                _CHAT_SESSIONS[session_id].append(
                    {"role": "assistant", "content": "".join(full_response)}
                )
            else:
                # LLM yielded zero tokens — remove the orphaned user message
                msgs = _CHAT_SESSIONS.get(session_id, [])
                if msgs and msgs[-1]["role"] == "user":
                    msgs.pop()
            yield "data: [DONE]\n\n"
        except Exception as e:
            print(f"[chat/stream] error: {e}")
            # Remove orphaned user message (no paired assistant response will be saved)
            msgs = _CHAT_SESSIONS.get(session_id, [])
            if msgs and msgs[-1]["role"] == "user":
                msgs.pop()
            yield f"data: {json.dumps({'error': str(e)})}\n\n"
            yield "data: [DONE]\n\n"

    return Response(
        generate(),
        content_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@app.route('/api/chat/session', methods=['GET'])
def get_chat_session():
    """Return stored messages for a session — used by frontend to restore on page reload."""
    session_id = request.args.get('session_id', '').strip()
    messages   = _CHAT_SESSIONS.get(session_id, [])
    return jsonify({"messages": messages})


@app.route('/api/chat/truncate', methods=['POST'])
def truncate_chat_session():
    """Keep only the first keep_up_to messages — used by Edit to discard subsequent turns."""
    data       = request.json or {}
    session_id = data.get('session_id', '').strip()
    try:
        keep_up_to = max(0, int(data.get('keep_up_to', 0)))
    except (ValueError, TypeError):
        return jsonify({"error": "keep_up_to must be a non-negative integer"}), 400
    if session_id in _CHAT_SESSIONS:
        _CHAT_SESSIONS[session_id] = _CHAT_SESSIONS[session_id][:keep_up_to]
    return jsonify({"success": True})


@app.route('/api/chat/stop', methods=['POST'])
def stop_chat_stream():
    """Client-initiated stop: remove orphaned user message if generation was aborted."""
    data       = request.json or {}
    session_id = data.get('session_id', '').strip()
    msgs       = _CHAT_SESSIONS.get(session_id, [])
    if msgs and msgs[-1]["role"] == "user":
        msgs.pop()
    return jsonify({"success": True})


@app.route('/api/chat/session', methods=['DELETE'])
def clear_chat_session():
    data       = request.json or {}
    session_id = data.get('session_id', '').strip()
    if session_id in _CHAT_SESSIONS:
        del _CHAT_SESSIONS[session_id]
    return jsonify({"success": True})


@app.route('/api/chat/download-docx', methods=['POST'])
def download_chat_docx():
    data       = request.json or {}
    session_id = data.get('session_id', '').strip()
    messages   = _CHAT_SESSIONS.get(session_id, [])

    if not messages:
        return jsonify({"error": "No conversation to download"}), 400

    doc = Document()
    doc.add_heading("Chat Conversation", 0)
    for msg in messages:
        p = doc.add_paragraph()
        p.add_run(f"{msg['role'].capitalize()}: ").bold = True
        p.add_run(msg['content'])
        doc.add_paragraph()

    tmp = tempfile.NamedTemporaryFile(suffix='.docx', delete=False, dir=UPLOAD_FOLDER)
    doc.save(tmp.name)
    tmp.close()
    return send_file(
        tmp.name,
        as_attachment=True,
        download_name='chat_conversation.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )


if __name__ == '__main__':
    os.makedirs('static', exist_ok=True)
    print("🚀 Agile Suite starting on http://localhost:5000")
    print("🔒 LOCAL MODE — all inference via llama-server, zero data leaves device")
    print("   Run: llama-server --model ~/models/qwen2.5-7b-instruct-q4_K_M.gguf \\")
    print("        --alias qwen2.5 --n-gpu-layers 999 --ctx-size 16384 --parallel 4 --port 8080")
    print(f"   Model: {LLM_MODEL} | Max concurrent: {_LLM_MAX_CONCURRENT}")
    if USE_CLOUD:
        print(f"⚠️  CLOUD OPT-IN ACTIVE — OpenRouter ({OPENROUTER_MODEL})")
        print("   PII_CLOUD_OPT_IN=1 set. Data WILL leave device. Unset to go fully local.")
    # Dev/single-user: threaded flask server. For production, use gunicorn:
    #   gunicorn --workers 1 --threads 16 --timeout 600 --worker-class gthread \
    #            --bind 0.0.0.0:5000 app:app
    #
    # MUST use --workers 1: _LLM_SEM and rate limiter are in-process — multiple
    # workers each get their own copy, breaking concurrency control entirely.
    # Single worker + 16 threads handles 50 concurrent users correctly.
    app.run(debug=False, port=5000, host='0.0.0.0', threaded=True)
